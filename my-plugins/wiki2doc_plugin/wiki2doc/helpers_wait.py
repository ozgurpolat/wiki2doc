""" Helper methods. """

import re
import docx
import urllib
from itertools import groupby
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
#from __builtin__ import None

FILTER_STYLES = [(r'(.*)\~\~(.*?)\~\~(.*)$',
                  r'(.*?)\~\~(.*?)\~\~', 'strike'),
                 (r'(.*)\,\,(.*?)\,\,(.*)$',
                  r'(.*?)\,\,(.*?)\,\,', 'subscript'),
                 (r'(.*)\^(.*?)\^(.*)$',
                  r'(.*?)\^(.*?)\^', 'superscript'),
                 (r'(.*)\*\*(.*?)\*\*(.*)$',
                  r'(.*?)\*\*(.*?)\*\*', 'bold'),
                 (r"(.*)\'\'\'(.*?)\'\'\'(.*)$",
                  r"(.*?)\'\'\'(.*?)\'\'\'", 'bold'),
                 (r"(.*)\'\'(.*?)\'\'(.*)$",
                  r"(.*?)\'\'(.*?)\'\'", 'italic'),
                 # there are two wiki formating statements for italic
                 # following causes error with hyper links because of
                 # //(.*? )//
                 #(r'(.*)\/\/(.*?)\/\/(.*)$', r'(.*?)\/\/(.*?)\/\/'),
                 (r'(.*)\*\*\'\'(.*?)\*\*\'\'(.*)$',
                  r'(.*?)\*\*\'\'(.*?)\*\*\'\'', 'bold', 'italic'),
                 (r"(.*)\'\'\'\'\'(.*?)\'\'\'\'\'(.*)$",
                  r"(.*?)\'\'\'\'\'(.*?)\'\'\'\'\'", 'bold', 'italic')]

FILTERS = [(r'^(.*)\\\\{2,}(\s*)(.*)$', r'(.*?)\\\\{2,}(\s*)'),
           (r"(.*)\{\{\{(.*?)\}\}\}(.*)$", r"(.*?)\{\{\{(.*?)\}\}\}"),
           (r'(.*)\[\s*\#(Fig\d+)\s*\](.*)$',
            r'(.*?)\[\s*\#(Fig\d+)\s*\]'),
           (r'(.*)\[\s*\#(Table\d+)\s*\](.*)$',
            r'(.*?)\[\s*\#(Table\d+)\s*\]'),
           (r'(.*)\[\s*(?:=|\s*)\#(Ref\d+)\s*\](.*)$',
            r'(.*?)\[\s*(?:=|\s*)\#(Ref\d+)\s*\]')]

def filter_wiki_text(text):
    """ for a given wiki text, this function filters
        the wiki text"""
    regex = r'(.*?)\\\\{1,}\s*$'
    text = filter_regex(regex, text)

    regex = r'^\s*=\s*(.*?)\s*=\s*$'
    text = filter_regex(regex, text)

    for flt in FILTERS:
        text = filter_multi_regex(flt[0], flt[1], text)

    return text

def filter_regex(regex, text):
    """ for a given wiki text, and regex, this function filters
        regex from the the wiki text"""
    regex_match = re.compile(regex)
    match = regex_match.match(text)
    if match:
        if match.group(1) == '':
            text = " "
        else:
            text = match.group(1)
    return text

def filter_multi_regex(regx, regy, text):
    """ for a given wiki text, and regex, this function filters
        regex repeated multiple times from the the wiki text """
    filter_match = re.compile(regx)
    match = filter_match.match(text)
    if match:
        matches = re.findall(regy,
                             text,
                             re.DOTALL)
        if len(matches) > 0:
            text = ''
            for each in matches:
                text += each[0]
                text += each[1]
            text += match.group(3)
    return text

def get_hyperlist_dbrk(hypermatches):
    """ Returns hyperlist for double
        bracked links"""

    hyperlist = []

    for hyper in hypermatches:
        linklist = re.split(r'\s+', hyper[2])
        hyperline = []
        spc = urllib.quote(' ')
        if len(linklist) > 1:
            hyperline.append(hyper[0])
            hyperline.append(hyper[1])
            path = ''
            for i in range(len(linklist)-1):
                path = path + linklist[i] + spc
            path = path + linklist[-1]
            path = path.strip()
            hyperline.append(path)
            hyperline.append(hyper[3])
            hyperline.append(hyper[4])
            hyperline = tuple(hyperline)
            hyperlist.append(hyperline)
        else:
            hyperlist.append(hyper)

    return hyperlist

def get_hyperlist_ticket(hypermatches):
    """ Returns hyperlist for ticket
        links (in Coconut tickets are
        task ids)"""

    hyperlist = []

    for hyper in hypermatches:
        hyperline = []
        hyperline.append(hyper[0])
        hyperline.append(hyper[1])
        hyperline.append(hyper[2])
        hyperline.append(hyper[3])
        hyperline.append(str(hyper[1]+hyper[2]))
        hyperline = tuple(hyperline)
        hyperlist.append(hyperline)

    return hyperlist

def check_for_relative_link(hypermatches):
    """ This function checks for ../
        and removes it. """

    hyperlist = []
    regex = re.compile(r'^\.\.\/(.*)')
    for hyper in hypermatches:
        if regex.match(hyper[1]):
            hyper = list(hyper)
            hyper[1] = str(regex.match(hyper[1]).group(1))
            hyper = tuple(hyper)
            hyperlist.append(hyper)
        else:
            hyperlist.append(hyper)
    return hyperlist


def find_hyperlinks(text):
    """ for a given text, this function finds multiple
        hyperlinks. There are three ways users can
        create hyperlinks:
        1. [[link_path|link_name]]
        2. [link_path link_name]
        3. e:wiki/link_path, wiki:/link_path
        plus:
        4. r:#{ID} reference to tickets"""

    regex = ''
    hyperlist = []
    rest = ''

    regex_id, hypermatches = select_link_type(text)
    #print('inside find_hyperlinks after -> regex_id, hypermatches = select_link_type(text)')
    #print('regex_id:', regex_id)
    #print('hypermatches:', hypermatches)

    if regex_id == 0 and len(hypermatches) > 0:
        #print('1. find_hyperlinks(text):')
        # matches [[link_path|link_name]]
        hyperlist = get_hyperlist_dbrk(hypermatches)
        regex = r'^(.*)\[\[(http\:|https\:|file\:|e\:\/wiki\/|e\:wiki\/|'+\
            r'\/wiki\/|wiki\:|attachment\:)(.*?)\]\](.*?)$'

    elif regex_id == 1 and len(hypermatches) > 0:
        #print('2. find_hyperlinks(text):')
        # matches [link_path link_name]
        hyperlist = hypermatches
        regex = r'^(.*)\[(http\:|https\:|file\:|e\:\/wiki\/|e\:wiki\/|' +\
            r'\/wiki\/|wiki\:|attachment\:)(.*?)\](.*?)$'

    elif regex_id == 2 and len(hypermatches) > 0:
        #print('3. find_hyperlinks(text):')
        hyperlist = hypermatches
        regex = r"^(.*)(http\:|https\:|file\:|e\:\/wiki\/|e\:wiki\/|" +\
            r"\/wiki\/|wiki\:|attachment\:)(.*?)(?:\s*$|\s+)(.*?)$"

    elif regex_id == 3 and len(hypermatches) > 0:
        #print('4. find_hyperlinks(text):')
        hyperlist = get_hyperlist_ticket(hypermatches)
        regex = r"^(.*)(r\:\#)(.*?)(?:\s*$|\s+)(.*?)$"

    elif regex_id == 4 and len(hypermatches) > 0:
        #print('5. find_hyperlinks(text):')
        hypermatches = check_for_relative_link(hypermatches)
        hyperlist = get_hyperlist_dbrk(hypermatches)
        regex = r'^(.*)\[\[(.*?\/.*?)(.*?)\]\](.*?)$'

    if regex_id >= 0 and len(hypermatches) > 0:
        #print('6. find_hyperlinks(text):')
        match_pattern = re.compile(regex)
        match = match_pattern.match(text)
        if match:
            rest = (match.group(4))
        if len(hyperlist) > 0:
            hyperlist.append(rest)
    return (regex_id, hyperlist)

def select_link_type(text):
    """ This function determines which
        regex matches the given text.. """

    regexes = [r"(.*?)" +\
               r"\[\[(http\:|https\:|file\:|e\:\/wiki\/|e\:wiki\/|" +\
               r"\/wiki\/|wiki\:|attachment\:)" +\
               r"(.*?)(\|(.*?)\]\]|\]\])",
               r"(.*?)" +\
               r"\[(http\:|https\:|file\:|e\:\/wiki\/|e\:wiki\/|" +\
               r"\/wiki\/|wiki\:|attachment\:)" +\
               r"(.*?)(\s+(.*?)\]|\])",
               r"(.*?)" +\
               r"(http\:|https\:|file\:|e\:\/wiki\/|e\:wiki\/|" +\
               r"\/wiki\/|wiki\:|attachment\:)" +\
               r"(.*?)((\s*$)|(\s+))",#(\s*$|\s+)(.*?)#((\s*)$|(\s+)(?:.*?))
               r"(.*?)" +\
               r"(r\:\#)" +\
               r"(.*?)((\s*$)|(\s+))",
               r"(.*?)" +\
               r"\[\[(.*?\/.*?)" +\
               r"(.*?)(\|(.*?)\]\]|\]\])"]#r:#805
    regex_id = 0
    hypermatches = []

    for idx, regex in enumerate(regexes):
        regex_id, hypermatches = get_hypermatches(idx, regex, text)
        if len(hypermatches) > 0:
            return (regex_id, hypermatches)

    return (None, [])

def get_hypermatches(idx, regex, text):
    """ This function finds all occurances of
        specified regex in a given text. """

    hypermatches = re.findall(regex,
                              text,
                              re.DOTALL)
    if len(hypermatches) > 0:
        return (idx, hypermatches)
    elif len(hypermatches) == 0:
        return (None, [])

def check_string(text):
    """ Checks to see if the string is
        unicode or ASCII in filter_hyperlinks """

    if isinstance(text, str):
        return unicode(text, "utf-8")
    elif isinstance(text, unicode):
        return text

def count_space(match):
    """ Counts number of spaces for a given
        text in filter_hyperlinks."""

    i = 0
    spc = ''
    while i < len(match):
        spc += ' '
        i += 1

    return spc

def process_blockquote(text):
    """ If there is space before and after the
        text HtmlFormatter creates blockquotes
        for this space which is later converted
        to linebreaks. This is not desirable.
        This methods deals with this problem. """

    regex = re.compile(r'(\s+)(.*)(\s+)')
    if regex.match(text):
        match_1 = regex.match(text).group(1)
        match_2 = regex.match(text).group(3)

        spc1 = count_space(match_1)
        spc2 = count_space(match_2)

        text = '{{{' + spc1 + '}}}' +\
            regex.match(text).group(2) +\
            '{{{' + spc2 + '}}}'

    return text

def table_font_size(table, size):
    """ this function sets table font size to 8Pt"""

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(size)
    return table

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file
    # and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url,
                          docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
                          is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rpr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        col = docx.oxml.shared.OxmlElement('w:color')
        col.set(docx.oxml.shared.qn('w:val'), color)
        rpr.append(col)

    # Set underlining
    und = docx.oxml.shared.OxmlElement('w:u')
    if underline:
        und.set(docx.oxml.shared.qn('w:val'), 'single')
    else:
        und.set(docx.oxml.shared.qn('w:val'), 'none')
    rpr.append(und)

    # Join all the xml elements together add add
    # the required text to the w:r element
    new_run.append(rpr)
    new_run.text = text
    hyperlink.append(new_run)

    # _p is protected and therefore not documented,
    # but otherwise it is impossible to add a hyperlink
    paragraph._p.append(hyperlink) # pylint: disable=protected-access

    return hyperlink

def insert_image(paragraph, img_path):
    """ insert image """

    if paragraph is not None:
        new = paragraph.insert_paragraph_before()
        run = new.add_run()
        run.add_picture(img_path, width=Inches(6.3))

def get_self_referencing_tasks(taskid_pairs):
    """ Parent child relationship of task ids are stored in mastertickets
        table in the database. It was discovered that some tickets (tasks)
        referenced itself in the mastertickets table. This causes function
        to enter into infinite loop. This method checks for circular refs
        in the mastertickets table and returns a list of ticket ids with
        circular references."""

    taskids = []
    for i in range(len(taskid_pairs)):
        if taskid_pairs[i][0] == taskid_pairs[i][1]:
            taskids.append(taskid_pairs[i][0])
    return taskids

def get_preceding_taskids(taskid, taskid_pairs):
    """ Parent child relationship of task ids are stored in mastertickets
        table in the database. This method receives a parent (successor)
        taskid and filtered task id pairs for the milestone and returns
        a list of all the preceding task ids for that task """

    taskids = []
    for i in range(len(taskid_pairs)):
        if taskid_pairs[i][1] == int(taskid):
            taskids.append(taskid_pairs[i][0])
    return taskids

def get_pre_ids_types_tasks(tasks):
    """ For a given list of tasks, returns a list of taskid,
        task type and the task info list stored in a tuple """

    taskids_type_tasks = []
    id_type_task = []

    for task in tasks:
        id_type_task = (
            task[TaskCols.ticket],
            task[TaskCols.type],
            task)
        taskids_type_tasks.append(id_type_task)

    taskids_type_tasks = sorted(taskids_type_tasks, key=lambda x: x[0])
    return taskids_type_tasks

def tables_in_spec_text(i_text):
    """ generator to iterate through the tables in the text of a spec,
    extracting the tables as lists of lists, removing them and placing
    anchors in their place. returns extracted tables and the text without
    the tables but with anchors in their place. """

    table = []
    regex = re.compile(r"^\s*\|\|")
    found = False
    key = 0
    text_without_tables = []
    
    lines = i_text[1].splitlines()
    print('lines', lines)
    for line in lines:
        match = regex.match(line)
        text_without_tables.append(line + '\n')
        print('line, match, found:', line, match, found)
        if match:
            found = True
            print('match line:', line)
            print('1.found', found)
            columns = line.split("||")
            
            for i in range(len(columns)):
                columns[i] = columns[i].strip()
            print('columns', columns)
            
            columns = columns[1:-1] #Removing first and last ||
            # columns = [list(j) for i, j in groupby(columns)]
            # The groupby(columns) statement is the short form
            # of the loop below, but due to pylint it is changed
            columnlist = []
            columnkeys = []
            for i, j in groupby(columns):
                columnlist.append(list(j))
                columnkeys.append(i)
            print('columnlist', columnlist)
            table.append(columnlist)
            text_without_tables.pop()
        elif found:
            print('found line:', line)
            print('2.found', found)
            # Inserting [[Table(Table_ID.tbl)]] anchor
            # and removing the table from the text
            line_after = text_without_tables[-1]
            
            print('line_after', line_after)
            print('text_without_tables', text_without_tables)
            
            text_without_tables.pop()
            key_string = 'Table_' + str(i_text[0]+1) + str(key+1)
            line = '[[Table(' + key_string + '.tbl)]]\n'
            text_without_tables.append(line)
            #text_without_tables.append(line_after + '\n')
            print('line_after', line_after)
            text_without_tables.append(line_after)
            yield (table, text_without_tables)
            found = False
            table = []
            key += 1
            text_without_tables = []
        else:
            found = False
        
    yield (table, text_without_tables)

def get_header_in_text_line(line):
    """ Find a header in a text line ....  """

    #header = re.compile(r'\s*(=+)\s*(\d*)')
    header = re.compile(r'\s*(=+)(\s*)(\.*)')

    match = header.match(line)
    
    print('get_header_in_text_line:')
    print('match', match)

    if match:
        print('match.group(1)', match.group(1))

    return line

def get_tables_in_text(sections):
    """ given a list of sections, returns a list of sections
        with attached tables stored in a dictionary where key
        is the table name in the spec and value is the table data. """

    sections_with_tables = []
    spec_tables = {}
    table_keys = []
    table_values = []
    key = 0
    text_without_tables = []
    for i in range(len(sections)):
        i_text = [i, sections[i][1]]
        if i_text[1] is not None:
            for table_text in tables_in_spec_text(i_text):
                if table_text and len(table_text[0]) > 0:
                    key_string = 'Table_' + str(i+1) + str(key+1)
                    table_keys.append(key_string)
                    table_values.append(table_text[0])
                    line = "".join(table_text[1])
                    text_without_tables.append(line)
                    key += 1
                else:
                    line = "".join(table_text[1])
                    text_without_tables.append(line)
        key = 0
        spec_tables = dict(zip(table_keys, table_values))
        text_without_tables = "\n".join(text_without_tables)
        spec_images = sections[i][2]
        sections_with_tables.append([sections[i][0],
                                     text_without_tables,
                                     spec_images,
                                     spec_tables])
        table_keys = []
        table_values = []
        text_without_tables = []
        spec_tables = {}

    return sections_with_tables

def merge_table(table, merge_list):
    """ For a given table data, and list of cells to be merged
       stored for each row of the table, it merges the cells, and
       returns a table with merged cells.
        Wiki Markup:
        || 1 || 2 || 3 ||
        |||| 1-2 || 3 ||
        || 1 |||| 2-3 ||
        |||||| 1-2-3 ||

        Display:
        ---- --- ----
        | 1 | 2 | 3 |
        ---- --- ----
        | 1-2   | 3 |
        -------- ----
        | 1 | 2-3   |
        -------------
        | 1-2-3     |
        -------------"""
    merge = merge_list
    for idx, row in enumerate(table.rows):
        for idy in range(len(merge[idx])):
            if len(merge[idx][idy]) > 0:
                cell_a = row.cells[merge[idx][idy][0]]
                cell_b = row.cells[merge[idx][idy][1]]
                merged_cells = [row.cells[merge[idx][idy][0]].text,
                                row.cells[merge[idx][idy][1]].text]
                cell_a.merge(cell_b)
                merged_cell_text = ''.join(merged_cells)
                row.cells[merge[idx][idy][0]].text = merged_cell_text
    return table

def get_spec_section(text, header, end):
    """ extract section with given header from spec """

    section = None
    start = re.compile(header)
    end = re.compile(end)
    found = False
    for line in text.splitlines():
        if found:
            if end.match(line):
                return section
            else:
                if section is None:
                    section = ""
                section += line + "\n"
                continue
        if start.match(line):
            found = True
    if found == False:
        return None
    else:
        return False

def get_storage_of_data(tasks):
    """ given a list of analyse apo task list, returns
        storage of data information. """

    intro_text = []
    section = []
    brackets = r"\[\[(.*)\]\]|\[(.*)\]"
    match_pattern = re.compile(brackets)
    link = ''
    path = ''
    pathname = ''

    for i in range(len(tasks)):
        match = match_pattern.match(tasks[i][TaskCols.storage_of_data])
        if match:
            if match.group(1):
                link = match.group(1).strip()
                splitlink = link.split('|')
            elif match.group(2):
                link = match.group(2).strip()
                splitlink = re.split(r'\s+', link)
            if len(splitlink) == 2:
                path = splitlink[0].strip()
                pathname = splitlink[1].strip()
                link = '[[' + path + '|' + pathname + ']]'
            else:
                path = link
                pathname = "Please Check for pipe symbol or spaces"
                link = '[[' + path + ' <- ' + pathname + ']]'
        text = 'Storage of Data: ' + link
        intro_text = [tasks[i][0], tasks[i][1], text]
        section.append(intro_text)
    return section

def request_redirect(req):
    """ Redirect request """

    if req.method == 'POST':
        if (req.args.get('project') is not None) and\
            (req.args.get('milestone') is not None) and\
            (req.args.get('igrmilestone') is not None) and\
            (req.args.get('get_igr_tasks') is not None) and\
            (req.args.get('get_tasks') is not None):
            link = '/autorep?project=' + \
                urllib.quote(req.args.get('project')) + \
                '&igrmilestone=' + \
                urllib.quote(req.args.get('igrmilestone')) + \
                '&milestone=' + \
                urllib.quote(req.args.get('milestone')) + \
                '&get_igr_tasks=' + \
                urllib.quote(req.args.get('get_igr_tasks')) + \
                '&get_tasks=' + \
                urllib.quote(req.args.get('get_tasks'))

            req.redirect(req.base_path + link)

        elif (req.args.get('project') is not None) and\
            (req.args.get('milestone') is not None) and\
            (req.args.get('igrtask') is not None) and\
            (req.args.get('task') is not None) and \
            (req.args.get('create_report') is not None):
            link = '/autorep?project=' + \
                urllib.quote(req.args.get('project')) + \
                '&milestone=' + \
                urllib.quote(req.args.get('milestone')) + \
                '&igrtask=' + \
                urllib.quote(req.args.get('igrtask')) + \
                '&task=' + \
                urllib.quote(req.args.get('task'))

            checkbox = re.compile(r'chk_\d*')

            for key in req.args.iterkeys():
                match = checkbox.match(key)
                if match:
                    link = link + '&' + key + '='
                    link = link + urllib.quote(req.args.get(key))

            link = link + '&create_report=' + \
                urllib.quote(req.args.get('create_report'))

            req.redirect(req.base_path + link)
        else:
            return False

        return True

def set_req_keys(req):
    """ Sets request keys"""

    create_report = None
    form_token = None
    get_wiki_link = None
     
    for key, value in req.args.iteritems():
        if key == 'create_report':
            create_report = value
        elif key == '__FORM_TOKEN':
            form_token = value
        elif key == 'get_wiki_link':
            get_wiki_link = value
             
    req_keys = [create_report,
                form_token,
                get_wiki_link]
    
    return req_keys

def set_tasks(tasks):
    """ Returns a list of igr or ogr tasks."""
    tasklist = []
    for task in tasks:
        tasksplit = task[1].split(',')
        taskid_key_value = (tasksplit[0], task)
        tasklist.append(taskid_key_value)
    tasklist = sorted(tasklist, key=itemgetter(0))
    tasks = []
    for task in tasklist:
        tasks.append(task[1])
    return tasks

def set_ttype(milestone, taskset, ttype_i, ttype_ii):
    """ Filters tasks for a given list of
        task type ([ttype_i, ttype_ii])."""
    tasks = set()
    for task in taskset:
        if task[2] == ttype_i or task[2] == ttype_ii:
            task_info = str(task[0]) + ', ' + \
                str(task[1].encode('utf-8')) + ', ' + \
                str(task[2])
            miles_task = (milestone,
                          to_unicode(task_info))
            tasks.add(miles_task)
    return tasks

def set_list_of_milestones(milestones):
    """ Filters milestones for duplicates and
        returns set of projects and milestones."""
    projects = set()
    set_of_milestones = set()
    for (project, milestone) in milestones:
        projects.add(project)
        set_of_milestones.add(milestone)
    return (projects, set_of_milestones)

def set_sel_apo_tasks(req):
    """ Set selected apo tasks"""
    sel_apo_tasks = []
    for key, value in req.args.iteritems():
        checkbox = re.compile(r'chk_\d*')
        match = checkbox.match(key)
        if match:
            sel_apo_tasks.append(value)
    return sel_apo_tasks

def get_sel_apo_task_ids(sel_apo_tasks):
    """ Select create apo task ids. """
    sel_apo_task_ids = []
    for i in range(len(sel_apo_tasks)):
        tasksplit = sel_apo_tasks[i].split(',')
        sel_apo_task_ids.append(tasksplit[0])
    return sel_apo_task_ids

def check_table_row_length(col_size, row_length):
    """ for a given column size and row length, this function
        checks to see if they are equal, if not raises value error."""
    if col_size != row_length:
        return False
    else:
        return True

def get_base_url(req):
    """ Returns base url from the request object. """
    base_url = req.base_url
    url = r"https?://(.*)?"
    url_match = re.compile(url)
    match = url_match.match(base_url)
    if match:
        base_url = "http://" + str(match.group(1)) + "/"
        return base_url

def remove_forward_slash(text):
    """ Removes forward slash from the text. """
    regex = re.compile(r'^\/(.*)')
    match = regex.match(text)
    if match:
        return match.group(1)
    else:
        return text

def create_list(paragraph):
    """ Create unordered list in docx """

    p_pr = paragraph._p.get_or_add_pPr() # pylint: disable=protected-access
    p_stylepr = p_pr.get_or_add_pStyle()
    p_style = OxmlElement('w:pStyle')
    p_style.set(qn('w:val'), str('List Paragraph'))
    p_stylepr.insert(0, p_style)

    num_pr = p_pr.get_or_add_numPr()
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(0))
    ilfo = OxmlElement('w:ilfo')
    ilfo.set(qn('w:val'), str(1))
    num_pr.insert(0, ilvl)
    num_pr.insert(0, ilfo)

    p_pr.insert(0, num_pr)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), "1.2cm")
    ind.set(qn('w:right'), "0cm")
    p_pr.insert(0, ind)

    return paragraph

def get_link_name(hyper):
    """ for a given hypermatch this function
        returns the link name."""

    link_name = ''

    if hyper[4] == '' and len(hyper) == 5:
        link_name = hyper[1] + hyper[2]
    elif hyper[4] == '' and len(hyper) == 6:
        spc = ''
        if re.match(r"(\s+)", hyper[5]):
            spc = re.match(r"(\s+)", hyper[5]).group(1)
        link_name = hyper[1] + hyper[2] + spc
    else:
        link_name = hyper[4]

    return link_name

def get_wiki_specname(spec_name, hyper):
    """ returns the wiki page name for another
        page that is under same parent path.
        """

    given_path = remove_forward_slash(hyper[1]) + hyper[2]
    given_path_list = given_path.split("/")
    spec_name_list = spec_name.split("/")

    list_index = []

    for i, item in enumerate(spec_name_list):
        if item in set(given_path_list):
            list_index.append(i)

    if len(list_index) > 0:
        spec_name_list = spec_name_list[:list_index[0]]
    elif len(list_index) == 0:
        spec_name_list = spec_name_list[:-1]

    mod_spec_name = ''

    for item in spec_name_list:
        mod_spec_name += item + "/"

    mod_spec_name = mod_spec_name + given_path

    return mod_spec_name
""" class to create a word document """

import re
import os
import tempfile
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
#from trac.mimeview import Context
from trac.web.chrome import web_context
from trac.util.text import to_unicode
from trac.wiki.formatter import HtmlFormatter
from HTMLParser import HTMLParser
from helpers import get_header_in_text_line
from helpers import filter_wiki_text
from helpers import find_hyperlinks
from helpers import process_blockquote
from helpers import check_string
from helpers import check_table_row_length
from helpers import merge_table
from helpers import table_font_size
from helpers import insert_image
from helpers import get_link_name
from helpers import add_hyperlink
from helpers import create_list
from helpers import get_base_url
from helpers import get_wiki_specname
from helpers import remove_forward_slash


from parser import DocumentHTMLParser

class Doc(object): # pylint: disable=too-many-public-methods
    """ class to create a document in MS Word """

    def __init__(self, args):
        self.document = Document(args[0])
        self.env = args[1]
        self.wiki2doc = args[2]
        self.req = args[3]
        
        self.add_hyper_link = True

#     def insert_paragraph_after(self, paragraph, text=None, style=None):
#         """Insert a new paragraph after the given paragraph."""
#         new_p = OxmlElement("w:p")
#         paragraph._p.addnext(new_p)
#         new_para = Paragraph(new_p, paragraph._parent)
#         if text:
#             new_para.add_run(text)
#         if style is not None:
#             new_para.style = style
#         return new_para
#   
    def add_document(self, sections):
        """ adds intoduction section/s in the spec/s to
            the introduction section in the document."""
 
        paragraph = self.get_paragraph_after_regex(r"")
 
        self.insert_section(paragraph, sections, 3)

    def get_paragraph_after_regex(self, regex):
         """ helper function to be used before insert_paragraph_before() """
   
         regex = re.compile(regex)
         found = False
         idx = 0
         while not found and idx < len(self.document.paragraphs):
             par = self.document.paragraphs[idx]
             match = regex.match(par.text)
             if match:
                 found = True
             else:
                 idx += 1
         idx += 1
         if found:
             if idx < len(self.document.paragraphs):
                 return self.document.paragraphs[idx]
         return self.document.add_paragraph()

    def insert_section(self, paragraph, sections, level):
        """ Given paragraph location and sections data,
            inserts section text, if found images and
            if found tables."""
 
        page_images = {}

        for i in range(len(sections)):

            text = sections[i][1]
            page_images.update(sections[i][2])
            style_key = 'Heading '+ str(level)
            #paragraph.insert_paragraph_before(apo_spec, style=style_key)
            if text is not None:
                params = [i, paragraph, sections, text, page_images]
                self.find_sections(params)

    def find_sections(self, params):
        """ Given paragraph location and sections data,
            inserts section text, if found images and
            if found tables.
            params = [i,
                      paragraph,
                      sections,
                      text,
                      page_images]"""
        img_filename = None
        img_path = None
        wiki_filter = \
            [
             re.compile(r'\s*\[\[Image\((.*(\.jpg|\.png|\.gif))\)\]\]\s*'),
             re.compile(r'\s*\[\[Table\((.*)\.tbl\)\]\]\s*'),
             re.compile(r'\s*(=+)(.+?)(=+)'),
             re.compile(r'\s*\[\s*=\#Table(\d+)\s*\]\s*'),
             re.compile(r'\s*\[\s*=\#Fig(\d+)\s*\]\s*'),
             re.compile(r'\s*\*\s*(.*)')]
#         image = re.compile(r'\s*\[\[Image\((.*(\.jpg|\.png|\.gif))\)\]\]\s*')
#         anchor = re.compile(r'\s*\[\[Table\((.*)\.tbl\)\]\]\s*')
#         section = re.compile(r'\s*(=+)\s*(\d+\.){1,}\d*(.*)')
#         tbl = re.compile(r'\s*\[\s*=\#Table(\d+)\s*\]\s*')
#         fig = re.compile(r'\s*\[\s*=\#Fig(\d+)\s*\]\s*')

        for line in params[3].splitlines():
            line = to_unicode(line)
#             img_match = wiki_filter[0].match(line)
#             anc_match = wiki_filter[1].match(line)
#             sec_match = wiki_filter[2].match(line)
#             tbl_match = wiki_filter[3].match(line)
#             fig_match = wiki_filter[4].match(line)
 
            get_header_in_text_line(line)
             
            if wiki_filter[0].match(line):

                img_filename = to_unicode(wiki_filter[0].match(line).group(1))
                img_filename_list = img_filename.split(':')
                
                # Handiling image from another page [[Image(wiki:Another_page:hello_world.jpg)]]
                if len(img_filename_list) == 3:

                    img_filename = img_filename_list[-1]
                
                for key, value in params[4].iteritems():

                    if key == img_filename:
                        img_path = value
                        # if you want to include the image name
                        # insert the code below
                        # params[1].insert_paragraph_before(line)
                        insert_image(params[1], img_path)
            elif wiki_filter[1].match(line):
                self.get_table(params,
                               to_unicode(wiki_filter[1].match(line).group(1)))
            elif wiki_filter[2].match(line):
                print('heading found', wiki_filter[2].match(line).group(1), wiki_filter[2].match(line).group(2).strip())
                style_key = 'Heading' +\
                            ' ' + \
                            str(len(wiki_filter[2].match(line).group(1)))
                
                params[1].insert_paragraph_before(\
                    to_unicode(wiki_filter[2].match(line).group(2).strip()),
                    style=style_key)
                 
                #params[1] = self.insert_paragraph_after(params[1], "Paragraph One And A Half.")
#                 if params[1] is not None:
#                     new = params[1].insert_paragraph_before()
#                     new.text = 'test'
#                     p = new._p
#                     p.addnext(new._p)
                     
                # IMPORTANT temoprarily commented out was in on 9.12.2020
#                 if params[1] is not None:
#                     new = params[1].insert_paragraph_before(to_unicode(wiki_filter[2].match(line).group(2).strip()), style=style_key)
#                     #new.text = 'test'
#                     run = new.add_run('test')
#                     run.font.subscript = True
#                     p = new._p
#                     p.addnext(new._p)                    
                 
            elif wiki_filter[3].match(line):
                line = 'Table' + ' ' + str(wiki_filter[3].match(line).group(1))
                line = to_unicode(line)
                params[1].insert_paragraph_before(line, style='Caption')
            elif wiki_filter[4].match(line):
                line = 'Figure' + ' ' + str(wiki_filter[4].match(line).group(1))
                line = to_unicode(line)
                params[1].insert_paragraph_before(line, style='Caption')
            elif wiki_filter[5].match(line):
                line = str(wiki_filter[5].match(line).group(1))
                line = to_unicode(line)
                paragraph = create_list(\
                    params[1].insert_paragraph_before(text=' ',
                                                      style='List Bullet'))
                line = filter_wiki_text(line)
                args = [None,
                        paragraph,
                        line,
                        params[2][params[0]][0]]
                # page_name -> params[2][params[0]][0] = sections[0]
                self.filter_hyperlinks(args)
            else:
                line = filter_wiki_text(line)
                args = [None,
                        params[1].insert_paragraph_before(),
                        line,
                        params[2][params[0]][0]]
                # page_name -> params[2][params[0]][0] = sections[0]
                self.filter_hyperlinks(args)

    def filter_hyperlinks(self, args):
        """ for a given paragraph text or a table text,
            this function filters the table text
            and returns the table data
            args = [table,
                    paragraph,
                    text,
                    page_name]"""
   
        #context = Context.from_request(self.req, 'wiki')
        context = web_context(self.req, 'wiki')
        regex_id, hypermatches = find_hyperlinks(args[2])
        hyperlink = ''
        if len(hypermatches) > 0:
                link_name = ''
                rest = hypermatches.pop()
                for hyper in hypermatches:
                    flt_text = ''
                    if self.add_hyper_link == True:
                        #args[2].rows[args[0]].cells[args[1]].\
                        #    paragraphs[0].add_run(hyper[0])
                        wiki = process_blockquote(check_string(hyper[0]))
                        self.parse_html(args, context, wiki)
                        hyperlink, link_name = self.get_hyperlink(args[3],
                                                                  regex_id,
                                                                  hyper)
                        if hyperlink == None:
                            break
                        add_hyperlink(args[1],
                                      hyperlink,
                                      link_name,
                                      '0000FF',
                                      True)
                    elif self.add_hyper_link == False:
                        flt_text = flt_text + hyper[0]
                        #args[2].rows[args[0]].cells[args[1]].\
                        #    paragraphs[0].add_run(flt_text)
                        wiki = process_blockquote(check_string(flt_text))
                        self.parse_html(args, context, wiki)
                #args[2].rows[args[0]].cells[args[1]].paragraphs[0].add_run(rest)
                wiki = process_blockquote(check_string(rest))
                self.parse_html(args, context, wiki)
        else:
            wiki = process_blockquote(check_string(args[2]))
            self.parse_html(args, context, wiki)

        return (args[0], hypermatches)

    def parse_html(self, args, context, wiki):
        """ Parse html string to docx
        args[1] = paragraph,
        context,
        wiki,
        args[3] = page_name"""
   
        try:
            html_code = HtmlFormatter(self.env,
                                      context,
                                      wiki).generate()
            DocumentHTMLParser(self.document, args[1], html_code)
            return html_code
        except AttributeError:
            self.wiki2doc.errorlog.append(
                ("HtmlFormatter could not parse" +\
                 " the following wikitext: {}".format(wiki),
                 get_base_url(self.req) + 'wiki/' + args[3]))

    def get_content(self):
         """ save docx to path """
         _, out = tempfile.mkstemp()
         self.save(out)
         with open(out) as filehndl:
             content = filehndl.read()
         os.unlink(out)

         return content

    def save(self, path):
        """ save docx to path """
        self.document.save(path)

    def get_table(self, params, match_group):
        """ Gets table information from
            list of sections and calls
            insert_table method.
            example_sections = [
            [2,
             'Page1',
             'consetetur sadipscing elitr, sed diam nonumy eirmod tempor\n' +\
             '[[Image(Image1.jpg)]]\n' +\
             'invidunt ut labore et dolore magna \n' +\
             'aliquyam erat, sed diam \n' +\
             '[[Image(Image2.jpg)]]\nvoluptua.\n',
             {'Image1.jpg': u'/tmp/trac-tempenv-WbQieJ/files/attachments/' +\
              'wiki/bdc/bdc726f49cd502d4306404b090a5ddd13bb7dc0e/98c78c01' +\
              'ccdb21a78fd4f561e980ccd4d3a5a685.jpg',
              'Image2.jpg': u'/tmp/trac-tempenv-WbQieJ/files/attachments/' +\
              'wiki/bdc/bdc726f49cd502d4306404b090a5ddd13bb7dc0e/e8385af6' +\
              'dfec928ba93ae7b6ccdc2c5f2fcb89f8.jpg'},
             {'Table_11': [[[' 1 '], [' 2 '], [' 3 ']],
                           [[''], [' 1-2 '], [' 3 ']],
                           [[' 1 '], [''], [' 2-3 ']],
                           [['', ''], [' 1-2-3 ']]]}],
               ...]]
            params = [i,
                      paragraph,
                      sections,
                      text,
                      page_images]"""
   
        if params[2][params[0]][3]:
            for value in params[2][params[0]][3]:
                if value == match_group:
                    table_data = params[2][params[0]][3][value]

                    self.insert_table(params[1],
                                      table_data,
                                      params[2][params[0]][0])

    def insert_table(self, paragraph, table_data, page_name):
        """ insert table """
   
        # ************************************* IMPORTANT *******
        # ADD THIS! KeyError: u"no style with name 'TableGrid'"
        '''
        Trac detected an internal error:
        KeyError: u"no style with name TableGrid
        There was an internal error in Trac. It is recommended that you notify your local Trac administrator with the information needed to reproduce the issue.
         
        To that end, you could  a ticket.
         
        The action that triggered the error was:
         
        POST: /wiki2doc
        TracGuide The Trac User and Administration Guide        
        '''
   
        table = self.append_table(table_data, page_name)
        table.style = 'Table Grid'
   
        if paragraph is not None:
            new = paragraph.insert_paragraph_before()
            # _p and _tbl are protected and therefore not documented,
            # but otherwise it is impossible to insert a table,
            # and it would only be possible to append it to the
            # end of the document
            new._p.addnext(table._tbl) # pylint: disable=protected-access

    def append_table(self, data, page_name):
        """ for a given table data, this function analyzes the table,
            looks for cells to be merged inside the text as described
            below. Creates the table with the values first, then calls
            merge_table method to merge the cells if they need to be merged,
            then returns the table.
   
            Wiki Markup:
            || 1 || 2 || 3 ||
            |||| 1-2 || 3 ||
            || 1 |||| 2-3 ||
            |||||| 1-2-3 ||
   
            Display:
            ---- --- ----
            | 1 | 2 | 3 |
            ---- --- ----
            | 1-2   | 3 |
            -------- ----
            | 1 | 2-3   |
            -------------
            | 1-2-3     |
            -------------"""
        merge_list = []
        col_size = 0
   
        for item in data[0]:
            col_size += len(item)
   
        table = self.document.add_table(rows=1, cols=col_size)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER # pylint: disable=no-member
        args = [data,
                table,
                col_size,
                page_name]
        table, merge_list = self.find_merged_cells(args)
   
        if len(data) == len(merge_list):
            table = merge_table(table, merge_list)
            table = table_font_size(table, 8)
            return table
        else:
            self.wiki2doc.errorlog.append(
                "Merge cell list length and table length does not match." +\
                "Please check the merged cells in: {} \n".format(data[0]),
                'None')

    def find_merged_cells(self, args):
        """ for a given table data, analyses
            the data and finds merged cells.
            args = [data,
                    table,
                    col_size,
                    page_name]"""    

        merge_list = []
        table_row_length = set()
        for idr, row in enumerate(args[0]):
            row_length = 0
   
            for item in row:
                row_length += len(item)
            params = [idr,
                      row,
                      table_row_length,
                      args[2],
                      row_length,
                      args[1],
                      args[3]]
             
            args[1], table_row_length, merge_row = self.get_merge_row(params)
   
            merge_list.append(merge_row)
            merge_row = []
            row_length = 0
   
            if idr < len(args[0])-1:
                args[1].add_row()
   
        if len(list(table_row_length)) > 0:
            page_path = get_base_url(self.req) +\
                'wiki/' + args[3]
            self.wiki2doc.errorlog.append((
                "There might be an extra pipe || in the wikitext of" +\
                " a table that needs to be removed. Number of" +\
                " columns in each row must match including merged" +\
                " cells! Check the following table with a:" +\
                " header: {}".format(args[0][0]),
                page_path))
        return (args[1], merge_list)

    def get_merge_row(self, params):
         """ for a given table data, analyses
             the data and finds merged cells.
             params = [idr, row, table_row_length,
                       col_size, row_length, table,
                       page_name]"""
   
         merge_row = []
         col = 0
         pos = 0
         start = 0
         end = 0

         for idx, item in enumerate(params[1]):
             for idy, value in enumerate(item):
                 if check_table_row_length(params[3],
                                           params[4]):
                     value = filter_wiki_text(value)
                     #args = [table, paragraph,
                     #        text, task_id, page_name]
                     args = [params[5],
                             params[5].rows[params[0]].cells[col].paragraphs[0],
                             value,
                             params[6]]
                     #args[2].rows[args[0]].cells[args[1]].paragraphs[0]
                     params[5], _ = self.filter_hyperlinks(args)
                 else:
                     params[2].add(False)
                 col += 1
                 start = pos
                 end = pos + len(item)
                 if idy == 0 and value == '' and idx < len(params[1])-1:
                     merge_row.append([start, end])
                 elif idy == 0 and value == '' and \
                     idx == len(params[1])-1 and len(item) > 1:
                     merge_row.append([start, end-1])
                 elif idy == 0 and col == params[4] - 1 and \
                     len(merge_row) == 0:
                     merge_row.append([])
             pos += len(item)
   
         return (params[5], params[2], merge_row)

    def get_hyperlink(self, page_name, regex_id, hyper):
        """ for a given hypermatch this function
            returns the hyperlink and the link name."""
   
        link_name = get_link_name(hyper)
        hyperlink = self.get_link_path(page_name, regex_id, hyper)
   
        return (hyperlink, link_name)

    def get_link_path(self, page_name, regex_id, hyper):
        """ for a given hypermatch this function
            returns the hyperlink path."""
   
        hyperlink = ''
   
        if regex_id == 4:
            another_child_page_name = get_wiki_specname(page_name, hyper)
            page = \
                self.wiki2doc.get_wikipage(
                    remove_forward_slash(another_child_page_name))
            if not page:
                self.errorlog_missing_page(another_child_page_name,
                                           page_name)
            hyperlink = self.get_wiki_hyperlink(page_name, hyper)
        elif hyper[1] == '/wiki/':
            page = self.wiki2doc.get_wikipage(hyper[2])
            if not page:
                self.errorlog_missing_page(hyper[2], page_name)
            hyperlink = get_base_url(self.req) +\
                remove_forward_slash(hyper[1]) +\
                hyper[2]
        elif hyper[1] == 'e:/wiki/':
            page = self.wiki2doc.get_wikipage(hyper[2])
            if not page:
                self.errorlog_missing_page(hyper[2], page_name)
            hyperlink = get_base_url(self.req) +\
            'wiki/' + hyper[2]
        elif hyper[1] == 'wiki:':
            page = \
                self.wiki2doc.get_wikipage(remove_forward_slash(hyper[2]))
            if not page:
                self.errorlog_missing_page(hyper[2], page_name)
            hyperlink = get_base_url(self.req) +\
                'wiki/' + hyper[2]
        else:
            hyperlink = hyper[1] + hyper[2]
   
        return hyperlink
   
    def errorlog_missing_page(self, missing_spec, in_spec):
        """ Errorlog a wiki page that
            does not exist. """
        missing_spec = remove_forward_slash(missing_spec)
        self.wiki2doc.errorlog.append(\
            ("Specified link does not exist. Please check the " +\
             "full path and the name of the following link:'{}']".format(\
                                                           missing_spec),
             get_base_url(self.req) + 'wiki/' + in_spec))
   
    def get_wiki_hyperlink(self, page_name, hyper):
        """ returns the wiki page hyperlink path for
            another page that is under same parent
            path. See regex_id 4. in find_hyperlinks.
   
            Example wikipage: http://example.org/wiki/AA/IDX001/Dummy-AA-Database/IDX001-AA-Page-ID
   
            Example reference from inside the link above.
            [[Dummy-AA-Database/GDL/Downloads| MS-GDL]]
            [[Dummy-AA-Database/GDL/Desktop| MBJ-GDL]]
            [[/GDL/Downloads| MS-GDL]]
            [[GDL/Downloads| MS-GDL]]
            [[/IDX001/Dummy-AA-Database/GDL/Downloads| MS-GDL]]
            [[IDX001/Dummy-AA-Database/GDL/Downloads| MS-GDL]]
            [[IDX001/Dummy-AA-Database/GDL/Downloads]]
   
            This works because both "IDX001-AA-Page-ID" and
            "GDL/Downloads" are under:
   
            http://example.org/wiki/AA/IDX001/Dummy-AA-Database/
            """
   
        given_path = remove_forward_slash(hyper[1]) + hyper[2]
        given_path_list = given_path.split("/")
        full_wiki_path = get_base_url(self.req) +\
            "wiki/" + page_name
        full_path_list = full_wiki_path.split("/")
        protocol = full_path_list[0]
        full_path_list = full_path_list[2:]
        list_index = []
        hyperlink = ''
   
        for i, item in enumerate(full_path_list):
            if item in set(given_path_list):
                list_index.append(i)
   
        if len(list_index) > 0:
            full_path_list = full_path_list[:list_index[0]]
        elif len(list_index) == 0:
            full_path_list = full_path_list[:-1]
   
        mod_full_path = ''
   
        for item in full_path_list:
            mod_full_path += item + "/"
   
        hyperlink = protocol + "//" + mod_full_path + given_path
   
        return hyperlink


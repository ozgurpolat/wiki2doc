""" Helper methods. """

import re
import docx
import urllib
from itertools import groupby
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
#from __builtin__ import None

FILTER_STYLES = [(r'(.*)\~\~(.*?)\~\~(.*)$',
                  r'(.*?)\~\~(.*?)\~\~', 'strike'),
                 (r'(.*)\,\,(.*?)\,\,(.*)$',
                  r'(.*?)\,\,(.*?)\,\,', 'subscript'),
                 (r'(.*)\^(.*?)\^(.*)$',
                  r'(.*?)\^(.*?)\^', 'superscript'),
                 (r'(.*)\*\*(.*?)\*\*(.*)$',
                  r'(.*?)\*\*(.*?)\*\*', 'bold'),
                 (r"(.*)\'\'\'(.*?)\'\'\'(.*)$",
                  r"(.*?)\'\'\'(.*?)\'\'\'", 'bold'),
                 (r"(.*)\'\'(.*?)\'\'(.*)$",
                  r"(.*?)\'\'(.*?)\'\'", 'italic'),
                 # there are two wiki formating statements for italic
                 # following causes error with hyper links because of
                 # //(.*? )//
                 #(r'(.*)\/\/(.*?)\/\/(.*)$', r'(.*?)\/\/(.*?)\/\/'),
                 (r'(.*)\*\*\'\'(.*?)\*\*\'\'(.*)$',
                  r'(.*?)\*\*\'\'(.*?)\*\*\'\'', 'bold', 'italic'),
                 (r"(.*)\'\'\'\'\'(.*?)\'\'\'\'\'(.*)$",
                  r"(.*?)\'\'\'\'\'(.*?)\'\'\'\'\'", 'bold', 'italic')]

FILTERS = [(r'^(.*)\\\\{2,}(\s*)(.*)$', r'(.*?)\\\\{2,}(\s*)'),
           (r"(.*)\{\{\{(.*?)\}\}\}(.*)$", r"(.*?)\{\{\{(.*?)\}\}\}"),
           (r'(.*)\[\s*\#(Fig\d+)\s*\](.*)$',
            r'(.*?)\[\s*\#(Fig\d+)\s*\]'),
           (r'(.*)\[\s*\#(Table\d+)\s*\](.*)$',
            r'(.*?)\[\s*\#(Table\d+)\s*\]'),
           (r'(.*)\[\s*(?:=|\s*)\#(Ref\d+)\s*\](.*)$',
            r'(.*?)\[\s*(?:=|\s*)\#(Ref\d+)\s*\]')]

def get_base_url(req):
    """ Returns base url from the request object. """
    base_url = req.base_url
    url = r"https?://(.*)?"
    url_match = re.compile(url)
    match = url_match.match(base_url)
    if match:
        base_url = "http://" + str(match.group(1)) + "/"
        return base_url
    
def set_req_keys(req):
    """ Sets request keys"""
 
    create_report = None
    form_token = None
    get_wiki_link = None
      
    for key, value in req.args.iteritems():
        if key == 'create_report':
            create_report = value
        elif key == '__FORM_TOKEN':
            form_token = value
        elif key == 'get_wiki_link':
            get_wiki_link = value
              
    req_keys = [create_report,
                form_token,
                get_wiki_link]
     
    return req_keys

def get_tables_in_text(sections):
    """ given a list of sections, returns a list of sections
        with attached tables stored in a dictionary where key
        is the table name in the spec and value is the table data. """

    sections_with_tables = []
    spec_tables = {}
    table_keys = []
    table_values = []
    key = 0
    text_without_tables = []
    for i in range(len(sections)):
        i_text = [i, sections[i][1]]
        if i_text[1] is not None:
            for table_text in tables_in_spec_text(i_text):
                if table_text and len(table_text[0]) > 0:
                    key_string = 'Table_' + str(i+1) + str(key+1)
                    table_keys.append(key_string)
                    table_values.append(table_text[0])
                    line = "".join(table_text[1])
                    text_without_tables.append(line)
                    key += 1
                else:
                    line = "".join(table_text[1])
                    text_without_tables.append(line)
        key = 0
        spec_tables = dict(zip(table_keys, table_values))
        text_without_tables = "\n".join(text_without_tables)
        spec_images = sections[i][2]
        sections_with_tables.append([sections[i][0],
                                     text_without_tables,
                                     spec_images,
                                     spec_tables])
        table_keys = []
        table_values = []
        text_without_tables = []
        spec_tables = {}

    return sections_with_tables

def tables_in_spec_text(i_text):
    """ generator to iterate through the tables in the text of a spec,
    extracting the tables as lists of lists, removing them and placing
    anchors in their place. returns extracted tables and the text without
    the tables but with anchors in their place. """

    table = []
    regex = re.compile(r"^\s*\|\|")
    found = False
    key = 0
    text_without_tables = []
    
    lines = i_text[1].splitlines()
    print('lines', lines)
    for line in lines:
        match = regex.match(line)
        text_without_tables.append(line + '\n')
        print('line, match, found:', line, match, found)
        if match:
            found = True
            print('match line:', line)
            print('1.found', found)
            columns = line.split("||")
            
            for i in range(len(columns)):
                columns[i] = columns[i].strip()
            print('columns', columns)
            
            columns = columns[1:-1] #Removing first and last ||
            # columns = [list(j) for i, j in groupby(columns)]
            # The groupby(columns) statement is the short form
            # of the loop below, but due to pylint it is changed
            columnlist = []
            columnkeys = []
            for i, j in groupby(columns):
                columnlist.append(list(j))
                columnkeys.append(i)
            print('columnlist', columnlist)
            table.append(columnlist)
            text_without_tables.pop()
        elif found:
            print('found line:', line)
            print('2.found', found)
            # Inserting [[Table(Table_ID.tbl)]] anchor
            # and removing the table from the text
            line_after = text_without_tables[-1]
            
            print('line_after', line_after)
            print('text_without_tables', text_without_tables)
            
            text_without_tables.pop()
            key_string = 'Table_' + str(i_text[0]+1) + str(key+1)
            line = '[[Table(' + key_string + '.tbl)]]\n'
            text_without_tables.append(line)
            #text_without_tables.append(line_after + '\n')
            print('line_after', line_after)
            text_without_tables.append(line_after)
            yield (table, text_without_tables)
            found = False
            table = []
            key += 1
            text_without_tables = []
        else:
            found = False
        
    yield (table, text_without_tables)
    
def get_header_in_text_line(line):
    """ Find a header in a text line ....  """

    #header = re.compile(r'\s*(=+)\s*(\d*)')
    header = re.compile(r'\s*(=+)(\s*)(\.*)')

    match = header.match(line)
    
    print('get_header_in_text_line:')
    print('match', match)

    if match:
        print('match.group(1)', match.group(1))

    return line

def filter_wiki_text(text):
    """ for a given wiki text, this function filters
        the wiki text"""
    regex = r'(.*?)\\\\{1,}\s*$'
    text = filter_regex(regex, text)

    regex = r'^\s*=\s*(.*?)\s*=\s*$'
    text = filter_regex(regex, text)

    for flt in FILTERS:
        text = filter_multi_regex(flt[0], flt[1], text)

    return text

def filter_regex(regex, text):
    """ for a given wiki text, and regex, this function filters
        regex from the the wiki text"""
    regex_match = re.compile(regex)
    match = regex_match.match(text)
    if match:
        if match.group(1) == '':
            text = " "
        else:
            text = match.group(1)

    return text

def filter_multi_regex(regx, regy, text):
    """ for a given wiki text, and regex, this function filters
        regex repeated multiple times from the the wiki text """
    filter_match = re.compile(regx)
    match = filter_match.match(text)
    if match:
        matches = re.findall(regy,
                             text,
                             re.DOTALL)
        if len(matches) > 0:
            text = ''
            for each in matches:
                text += each[0]
                text += each[1]
            text += match.group(3)

    return text

def find_hyperlinks(text):
    """ for a given text, this function finds multiple
        hyperlinks. There are three ways users can
        create hyperlinks:
        1. [[link_path|link_name]]
        2. [link_path link_name]
        3. e:wiki/link_path, wiki:/link_path
        plus:
        4. r:#{ID} reference to tickets"""

    regex = ''
    hyperlist = []
    rest = ''

    regex_id, hypermatches = select_link_type(text)
    #print('inside find_hyperlinks after -> regex_id, hypermatches = select_link_type(text)')
    #print('regex_id:', regex_id)
    #print('hypermatches:', hypermatches)

    if regex_id == 0 and len(hypermatches) > 0:
        #print('1. find_hyperlinks(text):')
        # matches [[link_path|link_name]]
        hyperlist = get_hyperlist_dbrk(hypermatches)
        regex = r'^(.*)\[\[(http\:|https\:|file\:|e\:\/wiki\/|e\:wiki\/|'+\
            r'\/wiki\/|wiki\:|attachment\:)(.*?)\]\](.*?)$'

    elif regex_id == 1 and len(hypermatches) > 0:
        #print('2. find_hyperlinks(text):')
        # matches [link_path link_name]
        hyperlist = hypermatches
        regex = r'^(.*)\[(http\:|https\:|file\:|e\:\/wiki\/|e\:wiki\/|' +\
            r'\/wiki\/|wiki\:|attachment\:)(.*?)\](.*?)$'

    elif regex_id == 2 and len(hypermatches) > 0:
        #print('3. find_hyperlinks(text):')
        hyperlist = hypermatches
        regex = r"^(.*)(http\:|https\:|file\:|e\:\/wiki\/|e\:wiki\/|" +\
            r"\/wiki\/|wiki\:|attachment\:)(.*?)(?:\s*$|\s+)(.*?)$"

    elif regex_id == 3 and len(hypermatches) > 0:
        #print('4. find_hyperlinks(text):')
        hyperlist = get_hyperlist_ticket(hypermatches)
        regex = r"^(.*)(r\:\#)(.*?)(?:\s*$|\s+)(.*?)$"

    elif regex_id == 4 and len(hypermatches) > 0:
        #print('5. find_hyperlinks(text):')
        hypermatches = check_for_relative_link(hypermatches)
        hyperlist = get_hyperlist_dbrk(hypermatches)
        regex = r'^(.*)\[\[(.*?\/.*?)(.*?)\]\](.*?)$'

    if regex_id >= 0 and len(hypermatches) > 0:
        #print('6. find_hyperlinks(text):')
        match_pattern = re.compile(regex)
        match = match_pattern.match(text)
        if match:
            rest = (match.group(4))
        if len(hyperlist) > 0:
            hyperlist.append(rest)

    return (regex_id, hyperlist)

def select_link_type(text):
    """ This function determines which
        regex matches the given text.. """

    regexes = [r"(.*?)" +\
               r"\[\[(http\:|https\:|file\:|e\:\/wiki\/|e\:wiki\/|" +\
               r"\/wiki\/|wiki\:|attachment\:)" +\
               r"(.*?)(\|(.*?)\]\]|\]\])",
               r"(.*?)" +\
               r"\[(http\:|https\:|file\:|e\:\/wiki\/|e\:wiki\/|" +\
               r"\/wiki\/|wiki\:|attachment\:)" +\
               r"(.*?)(\s+(.*?)\]|\])",
               r"(.*?)" +\
               r"(http\:|https\:|file\:|e\:\/wiki\/|e\:wiki\/|" +\
               r"\/wiki\/|wiki\:|attachment\:)" +\
               r"(.*?)((\s*$)|(\s+))",#(\s*$|\s+)(.*?)#((\s*)$|(\s+)(?:.*?))
               r"(.*?)" +\
               r"(r\:\#)" +\
               r"(.*?)((\s*$)|(\s+))",
               r"(.*?)" +\
               r"\[\[(.*?\/.*?)" +\
               r"(.*?)(\|(.*?)\]\]|\]\])"]#r:#805
    regex_id = 0
    hypermatches = []

    for idx, regex in enumerate(regexes):
        regex_id, hypermatches = get_hypermatches(idx, regex, text)
        if len(hypermatches) > 0:
            return (regex_id, hypermatches)

    return (None, [])

def get_hypermatches(idx, regex, text):
    """ This function finds all occurances of
        specified regex in a given text. """

    hypermatches = re.findall(regex,
                              text,
                              re.DOTALL)
    if len(hypermatches) > 0:
        return (idx, hypermatches)
    elif len(hypermatches) == 0:
        return (None, [])

def process_blockquote(text):
    """ If there is space before and after the
        text HtmlFormatter creates blockquotes
        for this space which is later converted
        to linebreaks. This is not desirable.
        This methods deals with this problem. """

    regex = re.compile(r'(\s+)(.*)(\s+)')
    if regex.match(text):
        match_1 = regex.match(text).group(1)
        match_2 = regex.match(text).group(3)

        spc1 = count_space(match_1)
        spc2 = count_space(match_2)

        text = '{{{' + spc1 + '}}}' +\
            regex.match(text).group(2) +\
            '{{{' + spc2 + '}}}'

    return text
    
def check_string(text):
    """ Checks to see if the string is
        unicode or ASCII in filter_hyperlinks """

    if isinstance(text, str):
        return unicode(text, "utf-8")
    elif isinstance(text, unicode):
        return text
    
def check_table_row_length(col_size, row_length):
    """ for a given column size and row length, this function
        checks to see if they are equal, if not raises value error."""
    if col_size != row_length:
        return False
    else:
        return True
    
def merge_table(table, merge_list):
    """ For a given table data, and list of cells to be merged
       stored for each row of the table, it merges the cells, and
       returns a table with merged cells.
        Wiki Markup:
        || 1 || 2 || 3 ||
        |||| 1-2 || 3 ||
        || 1 |||| 2-3 ||
        |||||| 1-2-3 ||

        Display:
        ---- --- ----
        | 1 | 2 | 3 |
        ---- --- ----
        | 1-2   | 3 |
        -------- ----
        | 1 | 2-3   |
        -------------
        | 1-2-3     |
        -------------"""
    merge = merge_list
    for idx, row in enumerate(table.rows):
        for idy in range(len(merge[idx])):
            if len(merge[idx][idy]) > 0:
                cell_a = row.cells[merge[idx][idy][0]]
                cell_b = row.cells[merge[idx][idy][1]]
                merged_cells = [row.cells[merge[idx][idy][0]].text,
                                row.cells[merge[idx][idy][1]].text]
                cell_a.merge(cell_b)
                merged_cell_text = ''.join(merged_cells)
                row.cells[merge[idx][idy][0]].text = merged_cell_text
    return table

def table_font_size(table, size):
    """ this function sets table font size to 8Pt"""

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(size)
    return table

def insert_image(paragraph, img_path):
    """ insert image """

    if paragraph is not None:
        new = paragraph.insert_paragraph_before()
        run = new.add_run()
        run.add_picture(img_path, width=Inches(6.3))

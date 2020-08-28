# pylint: disable=too-many-lines
# -*- coding: utf-8 -*-
"""
Unit tests for wiki2doc.py
"""

import unittest
import sys
import re
import os
sys.path.append(
    os.path.abspath(os.path.join(os.path.dirname(__file__), os.path.pardir)))
from docx import Document
from wiki2doc import WikiToDoc
from autorep.helpers import\
check_table_row_length, merge_table,\
table_font_size, add_hyperlink,\
filter_wiki_text, select_link_type,\
get_hyperlist_dbrk, get_hyperlist_ticket,\
create_list, get_link_name,\
get_link_name, get_wiki_specname,\
remove_forward_slash
from autorep.autorep import StructuralAnalysis
from autorep.autorep import AnalysisDataCompilation
from trac.test import EnvironmentStub, MockRequest
from simplemultiproject.environmentSetup\
import smpEnvironmentSetupParticipant
from datetime import datetime
from trac.util.datefmt import utc
from trac.wiki import WikiPage
from trac.util.text import to_unicode
from trac.mimeview import Context

def revert_schema(env):
    """ when we've finished, we have to manually
    revert the schema back to vanilla trac """
    with env.db_transaction as dbt:
        for table in ('smp_project', 'smp_milestone_project',
                      'smp_version_project', 'smp_component_project'):
            dbt("DROP TABLE IF EXISTS %s" % dbt.quote(table))
        dbt("DELETE FROM system WHERE name='simplemultiproject_version'")

def docx2txt(docxfile):
    ''' helper function to extract text from docx '''
    document = Document(docxfile)
    text = []
    for par in document.paragraphs:
        text.append(par.text)
    return text

def tables2txt(docxfile):
    ''' helper function to extract text from the tables in a docx '''
    document = Document(docxfile)
    out = []
    for table in document.tables:
        table_text = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text)
            table_text.append(row_text)
        out.append(table_text)
    return out

def _insert_wiki_pages(env, pages):
    """ insert wiki pages """
    time = datetime(2001, 1, 1, 1, 1, 1, 0, utc)
    for name, text in pages.iteritems():
        page = WikiPage(env)
        page.name = name
        page.text = text
        page.save('author', 'comment', '::1', time)

class ReportTestCase(unittest.TestCase):# pylint: disable=too-many-instance-attributes, too-many-public-methods
    """ Tests for the basic report api """

    def setUp(self):

        filename = 'in.docx'
        filename_adc = 'in_adc.docx'
        filename_sar = 'in_sar.docx'

        document = Document()
        document_adc = Document()
        document_sar = Document()

        for header in ['APOS ANALYSED',
                       'Introduction',
                       'Sub-Component Description',
                       'Main Drawings Reference and Modification Applicability',
                       'Material Properties',
                       'Fastener Properties',
                       'FEM Description (Global and Local)',
                       'Load Cases and associated Criteria',
                       'Main Hypotheses and Methods',
                       'Applicable factors',
                       'Documents',
                       'Software']:
            document.add_heading(header, level=1)
            document_adc.add_heading(header, level=1)
            document_sar.add_heading(header, level=1)

        document_adc.add_heading('Storage of Data', level=1)
        document_sar.add_heading('Abbreviations and Units', level=1)

        document.save(filename)
        document_adc.save(filename_adc)
        document_sar.save(filename_sar)

        self.gr_api = AutoRep(EnvironmentStub()) # pylint: disable=too-many-function-args

        for instance in self.gr_api.instances:
            self.gr_api.envs[instance] = EnvironmentStub(
                default_data=True,
                enable=["trac.*",
                        "simplemultiproject.*",
                        "mastertickets.*"])

            with self.gr_api.envs[instance].db_transaction as dbt:
                revert_schema(self.gr_api.envs[instance])
                smpEnvironmentSetupParticipant(
                    self.gr_api.envs[instance]).upgrade_environment(dbt)

        args = [filename,
                self.gr_api,
                'ADC',
                MockRequest(self.gr_api.envs['task']),
                self.gr_api.envs]
        self.report = AnalysisDataCompilation(args)

        args = [filename_adc,
                self.gr_api,
                'ADC',
                MockRequest(self.gr_api.envs['task']),
                self.gr_api.envs]
        self.report_adc = AnalysisDataCompilation(args)

        args = [filename_sar,
                self.gr_api,
                'SAR',
                MockRequest(self.gr_api.envs['task']),
                self.gr_api.envs]
        self.report_sar = StructuralAnalysis(args)

        self.section = [[2,
                         'Specname1',
                         '[[Table(Table_11.tbl)]]\n' +\
                         'consetetur sadipscing elitr,' +
                         ' sed diam nonumy eirmod tempor\n' +\
                         'invidunt ut labore et dolore magna \n' +\
                         '[http://lorem/ipsum/app/doc1 PLM Light]' +\
                         ' aliquyam erat, sed diam \n' +\
                         'voluptua.\n',
                         {},
                         {'Table_11': [[[' 1 '], [' 2 '], [' 3 ']],
                                       [[''], [' 1-2 '], [' 3 ']],
                                       [[' 1 '], [''], [' 2-3 ']],
                                       [['', ''], [' 1-2-3 ']]]}],
                        [3,
                         'Specname2',
                         '[[Table(Table_21.tbl)]]\n' +\
                         'duo dolores et ea rebum. Stet clita \n' +\
                         # pylint: disable=anomalous-backslash-in-string
                         '([file:///\\lorem\data$\Ipsum\PLM Link])' +\
                         ' kasd gubergren, no sea takimata \n' +\
                         'sanctus est Lorem ipsum dolor sit amet.\n',
                         {},
                         {'Table_21': [[[' 1 '], [' 2 '], [' 3 ']],
                                       [[''], [' 1-2 '], [' 3 ']],
                                       [[' 1 '], [''], [' 2-3 ']],
                                       [['', ''], [' 1-2-3 ']]]}]]
        self.tables = [[[u' 1 ', u' 2 ', u' 3 '],
                        [u' 1-2 ', u' 1-2 ', u' 3 '],
                        [u' 1 ', u' 2-3 ', u' 2-3 '],
                        [u' 1-2-3 ', u' 1-2-3 ', u' 1-2-3 ']],
                       [[u' 1 ', u' 2 ', u' 3 '],
                        [u' 1-2 ', u' 1-2 ', u' 3 '],
                        [u' 1 ', u' 2-3 ', u' 2-3 '],
                        [u' 1-2-3 ', u' 1-2-3 ', u' 1-2-3 ']]]

        self.out = 'out.docx'
        self.out_adc = 'out_adc.docx'
        self.out_sar = 'out_sar.docx'

    def tearDown(self):
        for fname in ["in.docx",
                      "in_adc.docx",
                      "in_sar.docx",
                      "out.docx",
                      "out_adc.docx",
                      "out_sar.docx"]:
            if os.path.isfile(fname):
                os.unlink(fname)

    def test_check_table_row_length(self):# pylint: disable=too-many-branches
        """ Test check_table_row_length."""

        data = [[[' 1 '], [' 2 '], [' 3 ']],
                [[''], [' 1-2 '], [' 3 ']],
                [[' 1 '], [''], [' 2-3 ']],
                [['', ''], [' 1-2-3 ']]]

        table_row_length = True
        col_size = 0
        unused_vars = []

        for item in data[0]:
            col_size += len(item)

        for idr, row in enumerate(data):
            row_length = 0

            for item in row:
                row_length += len(item)

            for idx, item in enumerate(row):
                for idy, value in enumerate(item):
                    unused = (idx, idy, idr, value)
                    unused_vars.append(unused) # for pylint
                    if check_table_row_length(col_size,
                                              row_length):
                        pass
                    else:
                        table_row_length = False

        self.assertEqual(table_row_length,
                         True,
                         "Value of table_row_length variable" +\
                         " must be True!")

        self.gr_api.errorlog = []

        # Introducing additional || pipe in the last
        # row of the table. check_table_row_length
        # method return False

        data = [[[' 1 '], [' 2 '], [' 3 ']],
                [[''], [' 1-2 '], [' 3 ']],
                [[' 1 '], [''], [' 2-3 ']],
                [['', '', ''], [' 1-2-3 ']]]

        table_row_length = set()
        col_size = 0

        for item in data[0]:
            col_size += len(item)

        for idr, row in enumerate(data):
            row_length = 0

            for item in row:
                row_length += len(item)

            for idx, item in enumerate(row):
                for idy, value in enumerate(item):
                    if check_table_row_length(col_size,
                                              row_length):
                        pass
                    else:
                        table_row_length.add(False)

        self.assertEqual(len(list(table_row_length)),
                         1,
                         "Length of table_row_length set" +\
                         " must be equal to 1!")

    def test_parse_html(self):
        """ Test parse_html."""

        pages = {
            'Specname1': """Lorem ipsum dolor sit amet,
consetetur sadipscing elitr, sed diam nonumy eirmod tempor
invidunt ut labore et dolore magna aliquyam erat, sed diam
voluptua."""}

        _insert_wiki_pages(self.gr_api.envs['event'], pages)

        html_code = ''
        # pylint: disable=trailing-whitespace
        wiki = """* '''bold''', 
   ''' triple quotes !''' 
   can be bold too if prefixed by ! ''', 
 * ''italic''
 * '''''bold italic''''' or ''italic and
   ''' italic bold ''' ''
 * __underline__
 * {{{monospace}}} or `monospace`
   (hence `{{{` or {{{`}}} quoting)
 * ~~strike-through~~
 * ^superscript^ 
 * ,,subscript,,
 * **also bold**, //italic as well//, 
   and **'' bold italic **'' //(since 0.12)//
"""

        paragraph = self.report.get_paragraph_after_regex(r"Introduction")

        exp_html_code = u'<ul><li><strong>bold</strong>,' +\
            ' \n</li></ul><blockquote>\n<p>\n<strong>' +\
            ' triple quotes \'\'\' \ncan be bold too' +\
            ' if prefixed by ! </strong>,' +\
            ' \n</p>\n</blockquote>\n<ul><li><em>italic</em>\n</li>' +\
            '<li><strong><em>bold italic</em></strong> or <em>italic' +\
            ' and\n<strong> italic bold </strong> </em>\n</li>' +\
            '<li><span class="underline">underline</span>\n</li>' +\
            '<li><tt>monospace</tt> or <tt>monospace</tt>' +\
            '\n(hence <tt>{{{</tt> or <tt>`</tt> quoting)\n' +\
            '</li><li><del>strike-through</del>\n</li>' +\
            '<li><sup>superscript</sup> \n</li><li><sub>subscript</sub>' +\
            '\n</li><li><strong>also bold</strong>,' +\
            ' <em>italic as well</em>, \nand <strong>' +\
            '<em> bold italic </em></strong><em></em>' +\
            ' <em>(since 0.12)</em>\n</li></ul>'

        context = Context.from_request(self.report.req, 'wiki')

        task_id = 2

        spec_name = 'Specname1'

        args = [None,
                paragraph,
                None,
                task_id,
                spec_name]

        html_code = self.report.parse_html(args, context, wiki)

#         Important:
#         print"Code below prints missing characters:"
#         print repr(html_code)

        self.report.document.save(self.out)
        self.assertEqual(html_code,
                         exp_html_code,
                         "html codes do not match!")

        self.report.document.save(self.out)

        doc = [u'APOS ANALYSED',
               u'Introduction',
               u"Sub-Component Description\nbold," +\
               " \n\n triple quotes ''' \ncan be" +\
               " bold too if prefixed by ! , \nitalic\nbold" +\
               " italic or italic and italic bold  \nunderline\n" +\
               "monospace or monospace(hence {{{ or ` quoting)\n" +\
               "strike-through\nsuperscript \nsubscript\nalso" +\
               " bold, italic as well, \nand  bold italic" +\
               "  (since 0.12)\n\n",
               u'Main Drawings Reference and Modification Applicability',
               u'Material Properties',
               u'Fastener Properties',
               u'FEM Description (Global and Local)',
               u'Load Cases and associated Criteria',
               u'Main Hypotheses and Methods',
               u'Applicable factors',
               u'Documents',
               u'Software']

        self.assertEqual(docx2txt(self.out),
                         doc,
                         "Documents do not match in " +\
                         "test_parse_html!")

    def test_add_hyperlink(self): # pylint: disable=too-many-locals
        """ Test add_hyperlink."""

        #[[' http://www.test.com '],

        pages = {
            'Specname1': """Lorem ipsum dolor sit amet,
consetetur sadipscing elitr, sed diam nonumy eirmod tempor
invidunt ut labore et dolore magna aliquyam erat, sed diam
voluptua."""}

        _insert_wiki_pages(self.gr_api.envs['event'], pages)

        paragraph = self.report.get_paragraph_after_regex(r"Introduction")

        hypermatches = [(' Lorem ipsum, dolor sit amet,',
                         '/wiki/',
                         'Specname1',
                         'Test Specification'),
                        (') consetetur [=#Fig8] sadipscing' +\
                         ' elitr, L, LT or ST* ',
                         'http',
                         '://lorem/ipsum/app/doc1',
                         'PLM Light')]

        for hyper in hypermatches:
            paragraph.insert_paragraph_before(hyper[0])
            if hyper[1] == '/wiki/':
                page = self.gr_api.get_wikipage(hyper[2])
                hyperlink = page.env.base_url + hyper[1] + hyper[2]
            else:
                hyperlink = hyper[1] + hyper[2]
            hyperlink = add_hyperlink(paragraph.insert_paragraph_before(),
                                      hyperlink,
                                      hyper[3],
                                      '0000FF',
                                      True)

            for item in hyperlink.itertext():
                self.assertEqual(item,
                                 hyper[3],
                                 "Hyperlink names do not match.")

    def test_filter_hyperlinks(self): # pylint: disable=too-many-locals
        """ Test filter_hyperlinks."""

        pages = {
            'Specname1': """Lorem ipsum dolor sit amet,
consetetur sadipscing elitr, sed diam nonumy eirmod tempor
invidunt ut labore et dolore magna aliquyam erat, sed diam
voluptua."""}

        _insert_wiki_pages(self.gr_api.envs['event'], pages)

        data = [[[' 1 '], [' 2 '], [' 3 ']],
                [[' http://www.test.com '],
                 [' [/wiki/Specname1 Specname1] '],
                 [' 3 ']],
                [[''], [' 1-2 '], [' 3 ']],
                [[' 1 '], [''], [' 2-3 ']],
                [['', ''], [' 1-2-3 ']]]
        task_id = 1
        spec_name = 'Specname 1'

        table_row_length = set()
        col_size = 0

        for item in data[0]:
            col_size += len(item)

        table = self.report.document.add_table(rows=1, cols=col_size)
        unused_vars = []
        for idr, row in enumerate(data):
            row_length = 0
            col = 0
            for item in row:
                row_length += len(item)

            for idx, item in enumerate(row):
                for idy, value in enumerate(item):
                    unused = (idx, idy)
                    unused_vars.append(unused) # for pylint
                    if check_table_row_length(col_size,
                                              row_length):
                        args = [table,
                                table.rows[idr].cells[col].paragraphs[0],
                                value,
                                task_id,
                                spec_name]
                        table, _ = \
                            self.report.filter_hyperlinks(args)
                    else:
                        table_row_length.add(False)
                    col += 1
            row_length = 0
            col = 0
            if idr < len(data)-1:
                table.add_row()

        paragraph = self.report.get_paragraph_after_regex(r"Introduction")

        if paragraph is not None:
            new = paragraph.insert_paragraph_before()
            # _p and _tbl are protected and therefore not documented,
            # but otherwise it is impossible to insert a table,
            # and it would only be possible to append it to the
            # end of the document
            new._p.addnext(table._tbl) # pylint: disable=protected-access

        self.report.document.save(self.out)

        expected_table = [[[u' 1 ', u' 2 ', u' 3 '],
                           [u' ', u'  ', u' 3 '],
                           [u'', u' 1-2 ', u' 3 '],
                           [u' 1 ', u'', u' 2-3 '],
                           [u'', u'', u' 1-2-3 ']]]

        self.assertEqual(tables2txt(self.out),
                         expected_table,
                         "Returned merge list and expected " +\
                         "merge list do not match!")

    def test_filter_hyperlinks_ii(self): # pylint: disable=too-many-locals
        """ Test filter_hyperlinks."""

        pages = {
            'Specname1': """consetetur sadipscing elitr,
sed diam nonumy eirmod tempor
[http://www.test.com] invidunt ut labore et dolore magna 
[/wiki/Specname1 Specname1] aliquyam erat, sed diam 
[e:/wiki/Specname1 Specname1] aliquyam erat, sed diam 
[wiki:/Specname1 Specname1] aliquyam erat, sed diam 
voluptua.""",
            'Specname2': """sed diam nonumy eirmod tempor
duo dolores et ea rebum. Stet clita 
[[http://www.test.com]] kasd gubergren, no sea takimata 
[[/wiki/Specname2|Specname2]] 
[[e:/wiki/Specname2|Specname2]] 
[[wiki:/Specname2|Specname2]] 
sanctus est Lorem ipsum dolor sit amet.""",
            'Specname3': """consetetur sadipscing elitr,
sed diam nonumy eirmod tempor
http://www.test.com invidunt ut labore et dolore magna 
/wiki/Specname1 Specname1 aliquyam erat, sed diam 
e:/wiki/Specname1 Specname1 aliquyam erat, sed diam 
wiki:/Specname1 Specname1 aliquyam erat, sed diam 
voluptua."""}

        _insert_wiki_pages(self.gr_api.envs['event'], pages)

        sections = [
            [2,
             'Specname1',
             'consetetur sadipscing elitr,\n' +\
             'sed diam nonumy eirmod tempor\n' +\
             '[http://www.test.com] invidunt ut labore et dolore magna \n' +\
             '[/wiki/Specname1 Specname1] aliquyam erat, sed diam \n' +\
             '[e:/wiki/Specname1 Specname1] aliquyam erat, sed diam \n' +\
             '[wiki:/Specname1 Specname1] aliquyam erat, sed diam \n' +\
             'voluptua.\n',
             {},
             {'Table_11': [[[' 1 '], [' 2 '], [' 3 ']],
                           [[''], [' 1-2 '], [' 3 ']],
                           [[' 1 '], [''], [' 2-3 ']],
                           [['', ''], [' 1-2-3 ']]]}],
            [3,
             'Specname2',
             #'[[Table(Table_21.tbl)]]\n' +\ Thros an error check this later!
             'sed diam nonumy eirmod tempor\n' +\
             'duo dolores et ea rebum. Stet clita \n' +\
             '[[http://www.test.com]] kasd gubergren, no sea takimata \n' +\
             '[[/wiki/Specname2|Specname2]] \n' +\
             '[[e:/wiki/Specname2|Specname2]] \n' +\
             '[[wiki:/Specname2|Specname2]] \n' +\
             'sanctus est Lorem ipsum dolor sit amet.\n',
             {},
             {'Table_21': [[[' 1 '], [' 2 '], [' 3 ']],
                           [[''], [' 1-2 '], [' 3 ']],
                           [[' 1 '], [''], [' 2-3 ']],
                           [['', ''], [' 1-2-3 ']]]}],
            [4,
             'Specname3',
             'consetetur sadipscing elitr,\n' +\
             'sed diam nonumy eirmod tempor\n' +\
             'http://www.test.com invidunt ut labore et dolore magna \n' +\
             '/wiki/Specname3 aliquyam erat, sed diam \n' +\
             'e:/wiki/Specname3 aliquyam erat, sed diam \n' +\
             'wiki:/Specname3 aliquyam erat, sed diam \n' +\
             'voluptua.\n',
             {},
             {'Table_11': [[[' 1 '], [' 2 '], [' 3 ']],
                           [[''], [' 1-2 '], [' 3 ']],
                           [[' 1 '], [''], [' 2-3 ']],
                           [['', ''], [' 1-2-3 ']]]}]]

        paragraph = self.report.get_paragraph_after_regex(r"Introduction")
        level = 3

        exp_hypermatches = [[(u'', u'http:', u'//www.test.com',
                              u']', u'')],
                            [(u'', u'/wiki/', u'Specname1',
                              u' Specname1]', u'Specname1')],
                            [(u'', u'e:/wiki/', u'Specname1',
                              u' Specname1]', u'Specname1')],
                            [(u'', u'wiki:', u'/Specname1',
                              u' Specname1]', u'Specname1')],
                            [(u'', u'http:', u'//www.test.com',
                              u']]', u'')],
                            [(u'', u'/wiki/', u'Specname2',
                              u'|Specname2]]', u'Specname2')],
                            [(u'', u'e:/wiki/', u'Specname2',
                              u'|Specname2]]', u'Specname2')],
                            [(u'', u'wiki:', u'/Specname2',
                              u'|Specname2]]', u'Specname2')],
                            # Notice there are five elements due to
                            # ((\s*$)|(\s+)) at the end of regex
                            [(u'', u'http:', u'//www.test.com',
                              u' ', u'', u' ')],
                            [(u'', u'/wiki/', u'Specname3',
                              u' ', u'', u' ')],
                            [(u'', u'e:/wiki/', u'Specname3',
                              u' ', u'', u' ')],
                            [(u'', u'wiki:', u'/Specname3',
                              u' ', u'', u' ')]]

        for i in range(len(sections)):
            text = sections[i][2]
            apo_spec = str(sections[i][0]) + ", " + sections[i][1]
            style_key = 'Heading '+ str(level)
            paragraph.insert_paragraph_before(apo_spec, style=style_key)
            if text is not None:
                for cnt, line in enumerate(text.splitlines()):
                    cnt += 1
                    line = to_unicode(line)
                    args = [None,
                            paragraph,
                            line,
                            sections[i][0],
                            sections[i][1]]
                    _, hypermatches = self.report.filter_hyperlinks(args)
                    if len(hypermatches) > 0:
                        self.assertEqual(exp_hypermatches[(cnt-3)+4*i],
                                         hypermatches,
                                         "Hypermatches do not match in " +\
                                         "test_filter_spec_hyperlinks!")

        self.report.document.save(self.out)

        doc = [u'APOS ANALYSED',
               u'Introduction',
               u'2, Specname1',
               u'3, Specname2',
               u'4, Specname3',
               u'Sub-Component Descriptionconsetetur' +\
               ' sadipscing elitr,sed diam nonumy' +\
               ' eirmod tempor invidunt ut labore' +\
               ' et dolore magna  aliquyam erat,' +\
               ' sed diam  aliquyam erat, sed diam' +\
               '  aliquyam erat, sed diam voluptua.sed' +\
               ' diam nonumy eirmod temporduo dolores' +\
               ' et ea rebum. Stet clita  kasd gubergren,' +\
               ' no sea takimata    sanctus est Lorem ipsum' +\
               ' dolor sit amet.consetetur sadipscing' +\
               ' elitr,sed diam nonumy eirmod temporinvidunt' +\
               ' ut labore et dolore magna aliquyam erat,' +\
               ' sed diam aliquyam erat, sed diam aliquyam' +\
               ' erat, sed diam voluptua.',
               u'Main Drawings Reference and Modification Applicability',
               u'Material Properties',
               u'Fastener Properties',
               u'FEM Description (Global and Local)',
               u'Load Cases and associated Criteria',
               u'Main Hypotheses and Methods',
               u'Applicable factors',
               u'Documents',
               u'Software']

        self.assertEqual(docx2txt(self.out),
                         doc,
                         "Documents do not match in " +\
                         "test_filter_spec_hyperlinks!")

    def test_filter_hyperlinks_iii(self): # pylint: disable=too-many-locals
        """ Test filter_hyperlinks."""

        pages = {
            'Specname4': """consetetur sadipscing elitr,
sed diam nonumy eirmod tempor
r:#805 invidunt ut labore et dolore magna 
r:#806 aliquyam erat, sed diam 
r:#807 aliquyam erat, sed diam 
r:#808 aliquyam erat, sed diam 
voluptua."""}

        _insert_wiki_pages(self.gr_api.envs['event'], pages)

        sections = [
            [2,
             'Specname1',
             'consetetur sadipscing elitr,\n' +\
             'sed diam nonumy eirmod tempor\n' +\
             'r:#805 invidunt ut labore et dolore magna \n' +\
             'r:#806 aliquyam erat, sed diam \n' +\
             'r:#807 aliquyam erat, sed diam \n' +\
             'r:#808 aliquyam erat, sed diam \n' +\
             'voluptua.\n',
             {},
             {'Table_11': [[[' 1 '], [' 2 '], [' 3 ']],
                           [[''], [' 1-2 '], [' 3 ']],
                           [[' 1 '], [''], [' 2-3 ']],
                           [['', ''], [' 1-2-3 ']]]}]]

        paragraph = self.report.get_paragraph_after_regex(r"Introduction")

        level = 3

        exp_hypermatches = [[(u'', u'r:#', u'805', u' ', 'r:#805')],
                            [(u'', u'r:#', u'806', u' ', 'r:#806')],
                            [(u'', u'r:#', u'807', u' ', 'r:#807')],
                            [(u'', u'r:#', u'808', u' ', 'r:#808')]]

        for i in range(len(sections)):
            text = sections[i][2]
            apo_spec = str(sections[i][0]) + ", " + sections[i][1]
            style_key = 'Heading '+ str(level)
            paragraph.insert_paragraph_before(apo_spec, style=style_key)
            if text is not None:
                for cnt, line in enumerate(text.splitlines()):
                    cnt += 1
                    line = to_unicode(line)
                    args = [None,
                            paragraph,
                            line,
                            sections[i][0],
                            sections[i][1]]
                    _, hypermatches = self.report.filter_hyperlinks(args)
                    if len(hypermatches) > 0:
                        self.assertEqual(exp_hypermatches[cnt-3],
                                         hypermatches,
                                         "Hypermatches do not match in " +\
                                         "test_filter_spec_hyperlinks!")

        self.report.document.save(self.out)

        doc = [u'APOS ANALYSED',
               u'Introduction',
               u'2, Specname1',
               u'Sub-Component Descriptionconsetetur sadipscing' +\
               ' elitr,sed diam nonumy eirmod temporinvidunt' +\
               ' ut labore et dolore magna aliquyam erat,' +\
               ' sed diam aliquyam erat, sed diam aliquyam' +\
               ' erat, sed diam voluptua.',
               u'Main Drawings Reference and Modification Applicability',
               u'Material Properties',
               u'Fastener Properties',
               u'FEM Description (Global and Local)',
               u'Load Cases and associated Criteria',
               u'Main Hypotheses and Methods',
               u'Applicable factors',
               u'Documents',
               u'Software']

        self.assertEqual(docx2txt(self.out),
                         doc,
                         "Documents do not match in " +\
                         "test_filter_spec_hyperlinks!")

    def test_select_link_type(self): # pylint: disable=too-many-statements
        """ Test select_link_type."""

        text_list = ["wiki:SED/IPSUM-2017-Dolore-Magna/DIAM2/Stet/SED" +\
                     "_DIAM2_Stet_CL7_EIRMOD_Elitr_Magna_5_Aliquyam_Erat",
                     "wiki:SED/IPSUM-2017-Dolore-Magna/DIAM2/Stet/SED" +\
                     "_DIAM2_Stet_CL7_EIRMOD_Elitr_Magna_5_Aliquyam_Erat test",
                     "wiki:/SED/IPSUM-2017-Dolore-Magna/DIAM2/Stet/SED" +\
                     "_DIAM2_Stet_CL7_EIRMOD_Elitr_Magna_5_Aliquyam_Erat",
                     "wiki:/SED/IPSUM-2017-Dolore-Magna/DIAM2/Stet/SED" +\
                     "_DIAM2_Stet_CL7_EIRMOD_Elitr_Magna_5_Aliquyam_Erat test",
                     "http://www.test.com e:wiki/SED/IPSUM-2017-Dolore-" +\
                     "Magna/DIAM2/Stet/SED_DIAM2_Stet_CL7_EIRMOD_Elitr_" +\
                     "Magna_5_Aliquyam_Erat ",
                     "[[e:wiki/SED/IPSUM-2017-Dolore-Magna/DIAM2/Stet/SED" +\
                     "_DIAM2_Stet_CL7_EIRMOD_Elitr_Magna_5_Aliquyam_Erat]]" +\
                     " test",
                     "[[e:wiki/SED/IPSUM-2017-Dolore-Magna/DIAM2/Stet/SED" +\
                     "_DIAM2_Stet_CL7_EIRMOD_Elitr_Magna_5_Aliquyam_Erat|" +\
                     "Link1]] test",
                     "[e:wiki/SED/IPSUM-2017-Dolore-Magna/DIAM2/Stet/SED" +\
                     "_DIAM2_Stet_CL7_EIRMOD_Elitr_Magna_5_Aliquyam_Erat" +\
                     " Link1] test",
                     "[http://www.test.com link1], [e:wiki/SED/IPSUM-2017" +\
                     "-Dolore-Magna/DIAM2/Stet/SED_DIAM2_Stet_CL7_EIRMOD_" +\
                     "Elitr_Magna_5_Aliquyam_Erat Link2] test",
                     "[[http://www.test.com|link1]], [[e:wiki/SED/IPSUM-" +\
                     "2017-Dolore-Magna/DIAM2/Stet/SED_DIAM2_Stet_CL7_" +\
                     "EIRMOD_Elitr_Magna_5_Aliquyam_Erat|Link2]] test",
                     "[[http://www.test.com]] test",
                     "e:wiki/SED/IPSUM-2017-Dolore-Magna/DIAM2/Stet/SED" +\
                     "_DIAM2_Stet_CL7_EIRMOD_Elitr_Magna_5_Aliquyam_Erat",
                     "r:#805 r:#806"]

        exp_hypermatches = [[('', 'wiki:', 'SED/IPSUM-2017-Dolore-Magna/' +\
                              'DIAM2/Stet/SED_DIAM2_Stet_CL7_EIRMOD_' +\
                              'Elitr_Magna_5_Aliquyam_Erat', '', '', '')],
                            [('', 'wiki:', 'SED/IPSUM-2017-Dolore-Magna/' +\
                              'DIAM2/Stet/SED_DIAM2_Stet_CL7_EIRMOD_' +\
                              'Elitr_Magna_5_Aliquyam_Erat', ' ', '', ' ')],
                            [('', 'wiki:', '/SED/IPSUM-2017-Dolore-Magna/' +\
                              'DIAM2/Stet/SED_DIAM2_Stet_CL7_EIRMOD_' +\
                              'Elitr_Magna_5_Aliquyam_Erat', '', '', '')],
                            [('', 'wiki:', '/SED/IPSUM-2017-Dolore-Magna/' +\
                              'DIAM2/Stet/SED_DIAM2_Stet_CL7_EIRMOD_' +\
                              'Elitr_Magna_5_Aliquyam_Erat', ' ', '', ' ')],
                            [('', 'http:', '//www.test.com', ' ', '', ' '),
                             ('', 'e:wiki/', 'SED/IPSUM-2017-Dolore-Magna/' +\
                              'DIAM2/Stet/SED_DIAM2_Stet_CL7_EIRMOD_' +\
                              'Elitr_Magna_5_Aliquyam_Erat', ' ', ' ', '')],
                            [('', 'e:wiki/', 'SED/IPSUM-2017-Dolore-Magna/' +\
                              'DIAM2/Stet/SED_DIAM2_Stet_CL7_EIRMOD_' +\
                              'Elitr_Magna_5_Aliquyam_Erat', ']]', '')],
                            [('', 'e:wiki/', 'SED/IPSUM-2017-Dolore-Magna/' +\
                              'DIAM2/Stet/SED_DIAM2_Stet_CL7_EIRMOD_' +\
                              'Elitr_Magna_5_Aliquyam_Erat',
                              '|Link1]]', 'Link1')],
                            [('', 'e:wiki/', 'SED/IPSUM-2017-Dolore-Magna/' +\
                              'DIAM2/Stet/SED_DIAM2_Stet_CL7_EIRMOD_' +\
                              'Elitr_Magna_5_Aliquyam_Erat',
                              ' Link1]', 'Link1')],
                            [('', 'http:', '//www.test.com',
                              ' link1]', 'link1'),
                             (', ', 'e:wiki/', 'SED/IPSUM-2017-Dolore-Magna/' +\
                              'DIAM2/Stet/SED_DIAM2_Stet_CL7_EIRMOD_' +\
                              'Elitr_Magna_5_Aliquyam_Erat',
                              ' Link2]', 'Link2')],
                            [('', 'http:', '//www.test.com',
                              '|link1]]', 'link1'),
                             (', ', 'e:wiki/', 'SED/IPSUM-2017-Dolore-Magna/' +\
                              'DIAM2/Stet/SED_DIAM2_Stet_CL7_EIRMOD_' +\
                              'Elitr_Magna_5_Aliquyam_Erat',
                              '|Link2]]', 'Link2')],
                            [('', 'http:', '//www.test.com', ']]', '')],
                            [('', 'e:wiki/', 'SED/IPSUM-2017-Dolore-Magna/' +\
                              'DIAM2/Stet/SED_DIAM2_Stet_CL7_EIRMOD_' +\
                              'Elitr_Magna_5_Aliquyam_Erat', '', '', '')],
                            [('', 'r:#', '805', ' ', '', ' '),
                             ('', 'r:#', '806', '', '', '')]]

        idx_list = [2, 2, 2, 2, 2, 0, 0, 1, 1, 0, 0, 2, 3]

        for i, text in enumerate(text_list):
            idx, hypermatches = select_link_type(text)

            self.assertEqual(idx_list[i],
                             idx,
                             "Returned hypermatches list and expected " +\
                             "hypermatches list do not match!")

            self.assertEqual(exp_hypermatches[i],
                             hypermatches,
                             "Returned hypermatches list and expected " +\
                             "hypermatches list do not match!")

        # TESTING file: links with spaces in them
        # pylint: disable=anomalous-backslash-in-string
        text = "Storage of Data: [[file://\\en.tp.firm\Diam$" +\
            "\Com\SED_Dolore\Magna-Aliquyam_PEDx\006" +\
            " lorem ipsum dolor sit amet\02 labore" +\
            " magna\P6\SED-Wizard\SED-SPEC-Dummy|Sed-Diam]]"

        idx, hypermatches = select_link_type(text)

        exp_hypermatches = [('Storage of Data: ',
                             'file:',
                             '//\\en.tp.firm\\Diam$\\Com\\SED_Dolore\\' +\
                             'Magna-Aliquyam_PEDx\x06 lorem ipsum dolor' +\
                             ' sit amet\x02 labore magna\\P6\\SED-Wizard\\' +\
                             'SED-SPEC-Dummy',
                             '|Sed-Diam]]',
                             'Sed-Diam')]

        self.assertEqual(exp_hypermatches,
                         hypermatches,
                         "Returned hypermatches list and expected " +\
                         "hypermatches list do not match!")

        # TESTING html: links with spaces in them
        text = "Storage of Data: [[http://www.test.com/x06 lorem" +\
            " ipsum dolorsit amet\x02 labore magna/sed/#ptc1/" +\
            "tcomp/infoPage?oid=VR%3Awt.doc.WTDocument%3A978821&u8=1|" +\
            "PLM Light]]"

        idx, hypermatches = select_link_type(text)

        exp_hypermatches = [('Storage of Data: ',
                             'http:',
                             '//www.test.com/x06 lorem ipsum' +\
                             ' dolorsit amet\x02 labore magna/' +\
                             'sed/#ptc1/tcomp/infoPage?oid=' +\
                             'VR%3Awt.doc.WTDocument%3A978821&u8=1',
                             '|PLM Light]]',
                             'PLM Light')]

        self.assertEqual(exp_hypermatches,
                         hypermatches,
                         "Returned hypermatches list and expected " +\
                         "hypermatches list do not match!")

        text = "Storage of Data: [[file://\\en.tp.firm\Diam$" +\
            "\Com\SED_Dolore\Magna-Aliquyam_PEDx\006" +\
            " lorem ipsum dolor sit amet\02 labore" +\
            " magna\P6\SED-Wizard\SED-SPEC-Dummy]]"

        idx, hypermatches = select_link_type(text)

        exp_hypermatches = [('Storage of Data: ',
                             'file:',
                             '//\\en.tp.firm\\Diam$\\Com\\SED_Dolore' +\
                             '\\Magna-Aliquyam_PEDx\x06 lorem ipsum' +\
                             ' dolor sit amet\x02 labore magna\\P6\\' +\
                             'SED-Wizard\\SED-SPEC-Dummy',
                             ']]',
                             '')]

        self.assertEqual(exp_hypermatches,
                         hypermatches,
                         "Returned hypermatches list and expected " +\
                         "hypermatches list do not match!")

        text = "Storage of Data:[[" +\
            r"file:///\\lorem\data$\Ipsum\PLM|Link" +\
             "]]"

        idx, hypermatches = select_link_type(text)

        exp_hypermatches = [('Storage of Data:',
                             'file:', '///\\\\lorem\\data$\\Ipsum\\PLM',
                             '|Link]]',
                             'Link')]

        self.assertEqual(exp_hypermatches,
                         hypermatches,
                         "Returned hypermatches list and expected " +\
                         "hypermatches list do not match!")

        text = "[[Dummy-APO-Database/GPD/Material_Strength| MS-GPD]]"

        idx, hypermatches = select_link_type(text)

        exp_hypermatches = [('',
                             'Dummy-APO-Database/',
                             'GPD/Material_Strength',
                             '| MS-GPD]]',
                             ' MS-GPD')]

        exp_regex_id = 4

        self.assertEqual(
            hypermatches,
            exp_hypermatches,
            "list of hypermatches do not match!")

        self.assertEqual(
            idx,
            exp_regex_id,
            "Expected regex id does not match!")

        text = "[[/GPD/Material_Strength| MS-GPD]]"

        idx, hypermatches = select_link_type(text)

        exp_hypermatches = [('',
                             '/',
                             'GPD/Material_Strength',
                             '| MS-GPD]]',
                             ' MS-GPD')]

        exp_regex_id = 4

        self.assertEqual(
            hypermatches,
            exp_hypermatches,
            "list of hypermatches do not match!")

        self.assertEqual(
            idx,
            exp_regex_id,
            "Expected regex id does not match!")

        text = "[[GPD/Material_Strength| MS-GPD]]"

        idx, hypermatches = select_link_type(text)

        exp_hypermatches = [('',
                             'GPD/',
                             'Material_Strength',
                             '| MS-GPD]]',
                             ' MS-GPD')]

        exp_regex_id = 4

        self.assertEqual(
            hypermatches,
            exp_hypermatches,
            "list of hypermatches do not match!")

        self.assertEqual(
            idx,
            exp_regex_id,
            "Expected regex id does not match!")

        text = "[[/IP006/Dummy-APO-Database/GPD/Material_Strength| MS-GPD]]"

        idx, hypermatches = select_link_type(text)

        exp_hypermatches = [('',
                             '/',
                             'IP006/Dummy-APO-Database/' +\
                             'GPD/Material_Strength',
                             '| MS-GPD]]',
                             ' MS-GPD')]

        exp_regex_id = 4

        self.assertEqual(
            hypermatches,
            exp_hypermatches,
            "list of hypermatches do not match!")

        self.assertEqual(
            idx,
            exp_regex_id,
            "Expected regex id does not match!")

        text = "[[IP006/Dummy-APO-Database/GPD/Material_Strength| MS-GPD]]"

        idx, hypermatches = select_link_type(text)

        exp_hypermatches = [('',
                             'IP006/',
                             'Dummy-APO-Database/GPD/' +\
                             'Material_Strength',
                             '| MS-GPD]]',
                             ' MS-GPD')]

        exp_regex_id = 4

        self.assertEqual(
            hypermatches,
            exp_hypermatches,
            "list of hypermatches do not match!")

        self.assertEqual(
            idx,
            exp_regex_id,
            "Expected regex id does not match!")

        text = "[[IP006/Dummy-APO-Database/GPD/Material_Strength]]"

        idx, hypermatches = select_link_type(text)

        exp_hypermatches = [('',
                             'IP006/',
                             'Dummy-APO-Database/' +\
                             'GPD/Material_Strength',
                             ']]',
                             '')]

        exp_regex_id = 4

        self.assertEqual(
            hypermatches,
            exp_hypermatches,
            "list of hypermatches do not match!")

        self.assertEqual(
            idx,
            exp_regex_id,
            "Expected regex id does not match!")

    def test_get_hyperlist_dbrk(self):
        """ Test get_hyperlist_dbrk."""

        hypermatches_list = [[('', 'e:wiki/', 'SED/IPSUM-2017-Dolore' +\
                               '-Magna/DIAM2/Stet/SED_DIAM2_Stet_CL7_' +\
                               'EIRMOD_Elitr_Magna_5_Aliquyam_Erat', ']]',
                               '')],
                             [('', 'e:wiki/', 'SED/IPSUM-2017-Dolore' +\
                               '-Magna/DIAM2/Stet/SED_DIAM2_Stet_CL7_' +\
                               'EIRMOD_Elitr_Magna_5_Aliquyam_Erat',
                               '|Link1]]',
                               'Link1')],
                             [('', 'http', '://www.test.com', '|link1]]',
                               'link1'),
                              (', ', 'e:wiki/', 'SED/IPSUM-2017-Dolore' +\
                               '-Magna/DIAM2/Stet/SED_DIAM2_Stet_CL7_' +\
                               'EIRMOD_Elitr_Magna_5_Aliquyam_Erat',
                               '|Link2]]',
                               'Link2')],
                             [('', 'http', '://www.test.com', ']]', '')]]

        exp_hyperlists = [[('', 'e:wiki/', 'SED/IPSUM-2017-Dolore-Magna/' +\
                            'DIAM2/Stet/SED_DIAM2_Stet_CL7_EIRMOD_Elitr' +\
                            '_Magna_5_Aliquyam_Erat', ']]', '')],
                          [('', 'e:wiki/', 'SED/IPSUM-2017-Dolore-Magna/' +\
                            'DIAM2/Stet/SED_DIAM2_Stet_CL7_EIRMOD_Elitr' +\
                            '_Magna_5_Aliquyam_Erat',
                            '|Link1]]',
                            'Link1')],
                          [('', 'http', '://www.test.com',
                            '|link1]]',
                            'link1'),
                           (', ', 'e:wiki/', 'SED/IPSUM-2017-Dolore-Magna/' +\
                            'DIAM2/Stet/SED_DIAM2_Stet_CL7_EIRMOD_Elitr' +\
                            '_Magna_5_Aliquyam_Erat',
                            '|Link2]]',
                            'Link2')],
                          [('', 'http', '://www.test.com', ']]', '')]]

        for i, hypermatches in enumerate(hypermatches_list):
            hyperlist = get_hyperlist_dbrk(hypermatches)

            self.assertEqual(exp_hyperlists[i],
                             hyperlist,
                             "Returned hyperlist and expected " +\
                             "hyperlist do not match!")

        hypermatches = [('Storage of Data: ',
                         'file:',
                         '//\\en.tp.firm\\Diam$\\Com\\SED_Dolore\\' +\
                         'Magna-Aliquyam_PEDx\x06 lorem ipsum dolor' +\
                         ' sit amet\x02 labore magna\\P6\\SED-Wizard\\' +\
                         'SED-SPEC-Dummy',
                         '|Sed-Diam]]',
                         'Sed-Diam')]

        hyperlist = get_hyperlist_dbrk(hypermatches)

        exp_hyperlist = [('Storage of Data: ',
                          'file:',
                          '//\\en.tp.firm\\Diam$\\Com\\SED_Dolore\\' +\
                          'Magna-Aliquyam_PEDx\x06%20lorem%20ipsum%' +\
                          '20dolor%20sit%20amet\x02%20labore%20magna\\' +\
                          'P6\\SED-Wizard\\SED-SPEC-Dummy',
                          '|Sed-Diam]]',
                          'Sed-Diam')]

        self.assertEqual(exp_hyperlist,
                         hyperlist,
                         "Returned hyperlist and expected " +\
                         "hyperlist do not match!")

        hypermatches = [('Storage of Data: ',
                         'file', ':///\\\\lorem\\data$\\Ipsum\\PLM',
                         '|Link]]',
                         'Link')]

        hyperlist = get_hyperlist_dbrk(hypermatches)

        exp_hyperlist = [('Storage of Data: ',
                          'file',
                          ':///\\\\lorem\\data$\\Ipsum\\PLM',
                          '|Link]]',
                          'Link')]

        self.assertEqual(exp_hyperlist,
                         hyperlist,
                         "Returned hyperlist and expected " +\
                         "hyperlist do not match!")

    def test_get_hyperlist_ticket(self):
        """ Test get_hyperlist_dbrk."""

        hypermatches = [('', 'r:#', '805', ' ', '', ' '),
                        ('', 'r:#', '806', '', '', '')]

        hyperlist = get_hyperlist_ticket(hypermatches)

        exp_hyperlists = [('', 'r:#', '805', ' ', 'r:#805'),
                          ('', 'r:#', '806', '', 'r:#806')]

        self.assertEqual(exp_hyperlists,
                         hyperlist,
                         "Returned hyperlist and expected " +\
                         "hyperlist do not match!")

    def test_errorlog_missing_page(self):
        """ Test errorlog_missing_page."""

        self.gr_api.errorlog = []
        task_id = 2
        missing_spec = 'SED_Diam'
        in_spec = 'Specname1'
        self.report.errorlog_missing_page(task_id, missing_spec, in_spec)

        errorlog = [("Specified link does not exist." +\
                     " Please check the full path and" +\
                     " the name of the following link:'SED_Diam']",
                     'http://example.org/Coconut/task/ticket/2',
                     'http://example.org/Coconut/event/wiki/Specname1')]

        self.assertEqual(self.gr_api.errorlog,
                         errorlog,
                         "Errorlogs do not match!")

    def test_get_link_name(self):
        """ Test get_link_name."""

        hypermatches = [(u'', u'http:', u'//www.test.com', u']', u''),
                        (u'', u'/wiki/', u'Specname1',
                         u' Specname1]', u'Specname1'),
                        (u'', u'e:/wiki/', u'Specname1',
                         u' Specname1]', u'Specname1'),
                        (u'', u'wiki:', u'/Specname1',
                         u' Specname1]', u'Specname1'),
                        (u'', u'http:', u'//www.test.com', u']]', u''),
                        (u'', u'/wiki/', u'Specname2',
                         u'|Specname2]]', u'Specname2'),
                        (u'', u'e:/wiki/', u'Specname2',
                         u'|Specname2]]', u'Specname2'),
                        (u'', u'wiki:', u'/Specname2',
                         u'|Specname2]]', u'Specname2'),
                        (u'', u'http:', u'//www.test.com', u' ', u'', u' '),
                        (u'', u'/wiki/', u'Specname3', u' ', u'', u' '),
                        (u'', u'e:/wiki/', u'Specname3', u' ', u'', u' '),
                        (u'', u'wiki:', u'/Specname3', u' ', u'', u' ')]

        exp_linknames = [u'http://www.test.com',
                         u'Specname1',
                         u'Specname1',
                         u'Specname1',
                         u'http://www.test.com',
                         u'Specname2',
                         u'Specname2',
                         u'Specname2',
                         u'http://www.test.com' + u' ',
                         u'/wiki/Specname3' + u' ',
                         u'e:/wiki/Specname3' + u' ',
                         u'wiki:/Specname3' + u' ']

        for i in range(len(hypermatches)):
            link_name = get_link_name(hypermatches[i])
            self.assertEqual(exp_linknames[i],
                             link_name,
                             "Returned hyperlist and expected " +\
                             "hyperlist do not match!")

        # TESTING Hyperlinks with spaces!
        self.gr_api.errorlog = []
        hypermatch = ('Storage of Data: ',
                      'file:',
                      '//\\en.tp.firm\\Diam$\\Com\\' +\
                      'SED_Dolore\\Magna-Aliquyam_PEDx' +\
                      '\x06%20lorem%20ipsum%20dolor%20sit' +\
                      '%20amet\x02%20labore%20magna\\P6\\' +\
                      'SED-Wizard\\SED-SPEC-Dummy',
                      '|Sed-Diam]]',
                      'Sed-Diam')

        link_name = get_link_name(hypermatch)

        exp_link_name = "Sed-Diam"

        self.assertEqual(exp_link_name,
                         link_name,
                         "Returned hyperlist and expected " +\
                         "hyperlist do not match!")

        self.gr_api.errorlog = []
        hypermatch = ('Storage of Data: ',
                      'file:',
                      '//\\en.tp.firm\\Diam$\\Com\\SED_Dolore' +\
                      '\\Magna-Aliquyam_PEDx\x06%20lorem%20' +\
                      'ipsum%20dolor%20sit%20amet\x02%20labore' +\
                      '%20magna\\P6\\SED-Wizard\\SED-SPEC-Dummy',
                      ']]',
                      '')
        link_name = get_link_name(hypermatch)
        # IMPORTANT! use repr when printing
        #print"hyperlink:{}".format(repr(hyperlink))
        exp_link_name = 'file://\\en.tp.firm\\Diam$\\Com\\SED_Dolore' +\
            '\\Magna-Aliquyam_PEDx\x06%20lorem%20ipsum' +\
            '%20dolor%20sit%20amet\x02%20labore%20magna' +\
            '\\P6\\SED-Wizard\\SED-SPEC-Dummy'

        self.assertEqual(exp_link_name,
                         link_name,
                         "Returned hyperlist and expected " +\
                         "hyperlist do not match!")

    def test_get_link_path(self):
        """ Test get_link_path."""

        task_ids = [2, 2, 2, 2,
                    3, 3, 3, 3,
                    4, 4, 4, 4]
        regex_ids = [1, 1, 1, 1,
                     0, 0, 0, 0,
                     2, 2, 2, 2]
        spec_names = ['Specname1', 'Specname1', 'Specname1', 'Specname1',
                      'Specname2', 'Specname2', 'Specname2', 'Specname2',
                      'Specname3', 'Specname3', 'Specname3', 'Specname3']
        hypermatches = [(u'', u'http:', u'//www.test.com', u']', u''),
                        (u'', u'/wiki/', u'Specname1',
                         u' Specname1]', u'Specname1'),
                        (u'', u'e:/wiki/', u'Specname1',
                         u' Specname1]', u'Specname1'),
                        (u'', u'wiki:', u'/Specname1',
                         u' Specname1]', u'Specname1'),
                        (u'', u'http:', u'//www.test.com', u']]', u''),
                        (u'', u'/wiki/', u'Specname2',
                         u'|Specname2]]', u'Specname2'),
                        (u'', u'e:/wiki/', u'Specname2',
                         u'|Specname2]]', u'Specname2'),
                        (u'', u'wiki:', u'/Specname2',
                         u'|Specname2]]', u'Specname2'),
                        (u'', u'http:', u'//www.test.com', u' ', u'', u' '),
                        (u'', u'/wiki/', u'Specname3', u' ', u'', u' '),
                        (u'', u'e:/wiki/', u'Specname3', u' ', u'', u' '),
                        (u'', u'wiki:', u'/Specname3', u' ', u'', u' ')]

        exp_hyperlinks = [u'http://www.test.com',
                          u'http://example.org/Coconut/event/wiki/Specname1',
                          u'http://example.org/Coconut/event/wiki/Specname1',
                          u'http://example.org/Coconut/event/wiki//Specname1',
                          u'http://www.test.com',
                          u'http://example.org/Coconut/event/wiki/Specname2',
                          u'http://example.org/Coconut/event/wiki/Specname2',
                          u'http://example.org/Coconut/event/wiki//Specname2',
                          u'http://www.test.com',
                          u'http://example.org/Coconut/event/wiki/Specname3',
                          u'http://example.org/Coconut/event/wiki/Specname3',
                          u'http://example.org/Coconut/event/wiki//Specname3']

        for i in range(len(hypermatches)):
            hyperlink = \
                self.report.get_link_path(task_ids[i],
                                          spec_names[i],
                                          regex_ids[i],
                                          hypermatches[i])
            self.assertEqual(exp_hyperlinks[i],
                             hyperlink,
                             "Returned hyperlink and expected " +\
                             "hyperlink do not match!")

        # TESTING Hyperlinks with spaces!
        self.gr_api.errorlog = []
        regex_id = 0
        task_id = 2
        spec_name = 'Specname1'
        hypermatch = ('Storage of Data: ',
                      'file:',
                      '//\\en.tp.firm\\Diam$\\Com\\' +\
                      'SED_Dolore\\Magna-Aliquyam_PEDx' +\
                      '\x06%20lorem%20ipsum%20dolor%20sit' +\
                      '%20amet\x02%20labore%20magna\\P6\\' +\
                      'SED-Wizard\\SED-SPEC-Dummy',
                      '|Sed-Diam]]',
                      'Sed-Diam')

        hyperlink = \
            self.report.get_link_path(task_id,
                                      spec_name,
                                      regex_id,
                                      hypermatch)
        # IMPORTANT! use repr when printing
        #print"hyperlink:{}".format(repr(hyperlink))
        exp_hyperlink = 'file://\\en.tp.firm\\Diam$\\Com\\SED_Dolore' +\
            '\\Magna-Aliquyam_PEDx\x06%20lorem%20ipsum' +\
            '%20dolor%20sit%20amet\x02%20labore%20magna' +\
            '\\P6\\SED-Wizard\\SED-SPEC-Dummy'

        self.assertEqual(exp_hyperlink,
                         hyperlink,
                         "Returned hyperlink and expected " +\
                         "hyperlink do not match!")

        self.gr_api.errorlog = []
        regex_id = 0
        task_id = 3
        spec_name = 'Specname2'
        hypermatch = ('Storage of Data: ',
                      'file:',
                      '//\\en.tp.firm\\Diam$\\Com\\SED_Dolore' +\
                      '\\Magna-Aliquyam_PEDx\x06%20lorem%20' +\
                      'ipsum%20dolor%20sit%20amet\x02%20labore' +\
                      '%20magna\\P6\\SED-Wizard\\SED-SPEC-Dummy',
                      ']]',
                      '')
        hyperlink = \
            self.report.get_link_path(task_id,
                                      spec_name,
                                      regex_id,
                                      hypermatch)
        # IMPORTANT! use repr when printing
        #print"hyperlink:{}".format(repr(hyperlink))
        exp_hyperlink = 'file://\\en.tp.firm\\Diam$\\Com\\SED_Dolore' +\
            '\\Magna-Aliquyam_PEDx\x06%20lorem%20ipsum' +\
            '%20dolor%20sit%20amet\x02%20labore%20magna' +\
            '\\P6\\SED-Wizard\\SED-SPEC-Dummy'

        self.assertEqual(exp_hyperlink,
                         hyperlink,
                         "Returned hyperlink and expected " +\
                         "hyperlink do not match!")

    def test_get_wiki_specname(self): # pylint: disable=too-many-statements
        """ Test get_wiki_specname."""

        self.gr_api.errorlog = []
        task_id = 2
        spec_name = 'APO/IP006/Dummy-APO-Database/' +\
            'IP006-APO-Spec-Sill'
        hypermatch = ('',
                      'Dummy-APO-Database/',
                      'GPD/Material_Strength',
                      '| MS-GPD]]',
                      ' MS-GPD')

        mod_spec_name = get_wiki_specname(spec_name,
                                          hypermatch)

        exp_spec_name = 'APO/IP006/Dummy-APO-Database/' +\
            'GPD/Material_Strength'

        self.assertEqual(exp_spec_name,
                         mod_spec_name,
                         "Returned spec and expected " +\
                         "spec names do not match!")

        # CHECKING TO SEE THE ERRORLOG,
        # WHEN PAGE DOES NOT EXIST!
        page = \
            self.gr_api.get_wikipage(
                remove_forward_slash(mod_spec_name))

        if not page:
            self.report.errorlog_missing_page(task_id,
                                              mod_spec_name,
                                              spec_name)

        errorlog = [("Specified link does not exist." +\
                     " Please check the full path and" +\
                     " the name of the following link:" +\
                     "'APO/IP006/Dummy-APO-Database/GPD/" +\
                     "Material_Strength']",
                     'http://example.org/Coconut/task/ticket/2',
                     'http://example.org/Coconut/event/wiki/' +\
                     'APO/IP006/Dummy-APO-Database/' +\
                     'IP006-APO-Spec-Sill')]

        self.assertEqual(self.gr_api.errorlog,
                         errorlog,
                         "Errorlogs do not match!")

        # NOW PAGE EXISTS. ERRORLOG SHOULD BE EMPTY!
        self.gr_api.errorlog = []
        pages = {
            'APO/IP006/Dummy-APO-Database/' +\
            'IP006-APO-Spec-Sill': """Lorem ipsum dolor sit amet,
consetetur sadipscing elitr, sed diam nonumy eirmod tempor
invidunt ut labore et dolore magna aliquyam erat, sed diam
voluptua."""}

        _insert_wiki_pages(self.gr_api.envs['event'], pages)

        self.assertEqual(self.gr_api.errorlog,
                         [],
                         "Errorlogs do not match!")

        self.gr_api.errorlog = []
        task_id = 2
        spec_name = 'APO/IP006/Dummy-APO-Database/' +\
            'IP006-APO-Spec-Sill'
        hypermatch = ('',
                      '/',
                      'GPD/Material_Strength',
                      '| MS-GPD]]',
                      ' MS-GPD')
        mod_spec_name = get_wiki_specname(spec_name,
                                          hypermatch)

        exp_spec_name = 'APO/IP006/Dummy-APO-Database/' +\
            'GPD/Material_Strength'

        self.assertEqual(exp_spec_name,
                         mod_spec_name,
                         "Returned spec and expected " +\
                         "spec names do not match!")

        self.gr_api.errorlog = []
        task_id = 2
        spec_name = 'APO/IP006/Dummy-APO-Database/' +\
            'IP006-APO-Spec-Sill'
        hypermatch = ('',
                      'GPD/',
                      'Material_Strength',
                      '| MS-GPD]]',
                      ' MS-GPD')
        mod_spec_name = get_wiki_specname(spec_name,
                                          hypermatch)

        exp_spec_name = 'APO/IP006/Dummy-APO-Database/' +\
            'GPD/Material_Strength'

        self.assertEqual(exp_spec_name,
                         mod_spec_name,
                         "Returned spec and expected " +\
                         "spec names do not match!")

        self.gr_api.errorlog = []
        task_id = 2
        spec_name = 'APO/IP006/Dummy-APO-Database/' +\
            'IP006-APO-Spec-Sill'
        hypermatch = ('',
                      '/',
                      'IP006/Dummy-APO-Database/GPD/' +\
                      'Material_Strength',
                      '| MS-GPD]]',
                      ' MS-GPD')
        mod_spec_name = get_wiki_specname(spec_name,
                                          hypermatch)

        exp_spec_name = 'APO/IP006/Dummy-APO-Database/' +\
            'GPD/Material_Strength'

        self.assertEqual(exp_spec_name,
                         mod_spec_name,
                         "Returned spec and expected " +\
                         "spec names do not match!")

        self.gr_api.errorlog = []
        task_id = 2
        spec_name = 'APO/IP006/Dummy-APO-Database/' +\
            'IP006-APO-Spec-Sill'
        hypermatch = ('',
                      'IP006/',
                      'Dummy-APO-Database/GPD/' +\
                      'Material_Strength',
                      '| MS-GPD]]',
                      ' MS-GPD')
        mod_spec_name = get_wiki_specname(spec_name,
                                          hypermatch)

        exp_spec_name = 'APO/IP006/Dummy-APO-Database/' +\
            'GPD/Material_Strength'

        self.assertEqual(exp_spec_name,
                         mod_spec_name,
                         "Returned spec and expected " +\
                         "spec names do not match!")

        self.gr_api.errorlog = []
        task_id = 2
        spec_name = 'APO/IP006/Dummy-APO-Database/' +\
            'IP006-APO-Spec-Sill'
        hypermatch = ('',
                      'IP006/',
                      'Dummy-APO-Database/GPD/' +\
                      'Material_Strength',
                      ']]',
                      '')
        mod_spec_name = get_wiki_specname(spec_name,
                                          hypermatch)

        exp_spec_name = 'APO/IP006/Dummy-APO-Database/' +\
            'GPD/Material_Strength'

        self.assertEqual(exp_spec_name,
                         mod_spec_name,
                         "Returned spec and expected " +\
                         "spec names do not match!")

    def test_get_wiki_hyperlink(self):
        """ Test get_wiki_hyperlink."""

        self.gr_api.errorlog = []
        spec_name = 'APO/IP006/Dummy-APO-Database/' +\
            'IP006-APO-Spec-Sill'

        hypermatch = ('',
                      'Dummy-APO-Database/',
                      'GPD/Material_Strength',
                      '| MS-GPD]]',
                      ' MS-GPD')

        hyperlink = self.report.get_wiki_hyperlink(spec_name,
                                                   hypermatch)

        exp_hyperlink = "http://example.org/Coconut/" +\
            "event/wiki/APO/IP006/Dummy-APO-Database/" +\
            "GPD/Material_Strength"

        self.assertEqual(exp_hyperlink,
                         hyperlink,
                         "Returned and expected " +\
                         "hyperlinks do not match!")

        self.gr_api.errorlog = []
        spec_name = 'APO/IP006/Dummy-APO-Database/' +\
            'IP006-APO-Spec-Sill'
        hypermatch = ('',
                      '/',
                      'GPD/Material_Strength',
                      '| MS-GPD]]',
                      ' MS-GPD')

        hyperlink = self.report.get_wiki_hyperlink(spec_name,
                                                   hypermatch)

        exp_hyperlink = "http://example.org/Coconut/" +\
            "event/wiki/APO/IP006/Dummy-APO-Database/" +\
            "GPD/Material_Strength"

        self.assertEqual(exp_hyperlink,
                         hyperlink,
                         "Returned and expected " +\
                         "hyperlinks do not match!")

        self.gr_api.errorlog = []
        spec_name = 'APO/IP006/Dummy-APO-Database/' +\
            'IP006-APO-Spec-Sill'
        hypermatch = ('',
                      'GPD/',
                      'Material_Strength',
                      '| MS-GPD]]',
                      ' MS-GPD')

        hyperlink = self.report.get_wiki_hyperlink(spec_name,
                                                   hypermatch)

        exp_hyperlink = "http://example.org/Coconut/" +\
            "event/wiki/APO/IP006/Dummy-APO-Database/" +\
            "GPD/Material_Strength"

        self.assertEqual(exp_hyperlink,
                         hyperlink,
                         "Returned and expected " +\
                         "hyperlinks do not match!")

        self.gr_api.errorlog = []
        spec_name = 'APO/IP006/Dummy-APO-Database/' +\
            'IP006-APO-Spec-Sill'
        hypermatch = ('',
                      '/',
                      'IP006/Dummy-APO-Database/GPD/' +\
                      'Material_Strength',
                      '| MS-GPD]]',
                      ' MS-GPD')

        hyperlink = self.report.get_wiki_hyperlink(spec_name,
                                                   hypermatch)

        exp_hyperlink = "http://example.org/Coconut/" +\
            "event/wiki/APO/IP006/Dummy-APO-Database/" +\
            "GPD/Material_Strength"

        self.assertEqual(exp_hyperlink,
                         hyperlink,
                         "Returned and expected " +\
                         "hyperlinks do not match!")

        self.gr_api.errorlog = []
        spec_name = 'APO/IP006/Dummy-APO-Database/' +\
            'IP006-APO-Spec-Sill'
        hypermatch = ('',
                      'IP006/',
                      'Dummy-APO-Database/GPD/' +\
                      'Material_Strength',
                      '| MS-GPD]]',
                      ' MS-GPD')

        hyperlink = self.report.get_wiki_hyperlink(spec_name,
                                                   hypermatch)

        exp_hyperlink = "http://example.org/Coconut/" +\
            "event/wiki/APO/IP006/Dummy-APO-Database/" +\
            "GPD/Material_Strength"

        self.assertEqual(exp_hyperlink,
                         hyperlink,
                         "Returned and expected " +\
                         "hyperlinks do not match!")

        self.gr_api.errorlog = []
        spec_name = 'APO/IP006/Dummy-APO-Database/' +\
            'IP006-APO-Spec-Sill'
        hypermatch = ('',
                      'IP006/',
                      'Dummy-APO-Database/GPD/' +\
                      'Material_Strength',
                      ']]',
                      '')

        hyperlink = self.report.get_wiki_hyperlink(spec_name,
                                                   hypermatch)

        exp_hyperlink = "http://example.org/Coconut/" +\
            "event/wiki/APO/IP006/Dummy-APO-Database/" +\
            "GPD/Material_Strength"

        self.assertEqual(exp_hyperlink,
                         hyperlink,
                         "Returned and expected " +\
                         "hyperlinks do not match!")

    def test_get_hyperlink(self): # pylint: disable=too-many-locals
        """ Test get_hyperlink."""

        task_ids = [2, 2, 2, 2,
                    3, 3, 3, 3,
                    4, 4, 4, 4]
        regex_ids = [1, 1, 1, 1,
                     0, 0, 0, 0,
                     2, 2, 2, 2]
        spec_names = ['Specname1', 'Specname1', 'Specname1', 'Specname1',
                      'Specname2', 'Specname2', 'Specname2', 'Specname2',
                      'Specname3', 'Specname3', 'Specname3', 'Specname3']
        hypermatches = [(u'', u'http:', u'//www.test.com', u']', u''),
                        (u'', u'/wiki/', u'Specname1',
                         u' Specname1]', u'Specname1'),
                        (u'', u'e:/wiki/', u'Specname1',
                         u' Specname1]', u'Specname1'),
                        (u'', u'wiki:', u'/Specname1',
                         u' Specname1]', u'Specname1'),
                        (u'', u'http:', u'//www.test.com', u']]', u''),
                        (u'', u'/wiki/', u'Specname2',
                         u'|Specname2]]', u'Specname2'),
                        (u'', u'e:/wiki/', u'Specname2',
                         u'|Specname2]]', u'Specname2'),
                        (u'', u'wiki:', u'/Specname2',
                         u'|Specname2]]', u'Specname2'),
                        (u'', u'http:', u'//www.test.com', u' ', u'', u' '),
                        (u'', u'/wiki/', u'Specname3', u' ', u'', u' '),
                        (u'', u'e:/wiki/', u'Specname3', u' ', u'', u' '),
                        (u'', u'wiki:', u'/Specname3', u' ', u'', u' ')]

        exp_hyperlinks = [u'http://www.test.com',
                          u'http://example.org/Coconut/event/wiki/Specname1',
                          u'http://example.org/Coconut/event/wiki/Specname1',
                          u'http://example.org/Coconut/event/wiki//Specname1',
                          u'http://www.test.com',
                          u'http://example.org/Coconut/event/wiki/Specname2',
                          u'http://example.org/Coconut/event/wiki/Specname2',
                          u'http://example.org/Coconut/event/wiki//Specname2',
                          u'http://www.test.com',
                          u'http://example.org/Coconut/event/wiki/Specname3',
                          u'http://example.org/Coconut/event/wiki/Specname3',
                          u'http://example.org/Coconut/event/wiki//Specname3']

        exp_linknames = [u'http://www.test.com',
                         u'Specname1',
                         u'Specname1',
                         u'Specname1',
                         u'http://www.test.com',
                         u'Specname2',
                         u'Specname2',
                         u'Specname2',
                         u'http://www.test.com' + u' ',
                         u'/wiki/Specname3' + u' ',
                         u'e:/wiki/Specname3' + u' ',
                         u'wiki:/Specname3' + u' ']

        for i in range(len(hypermatches)):
            hyperlink, link_name = \
                self.report.get_hyperlink(task_ids[i],
                                          spec_names[i],
                                          regex_ids[i],
                                          hypermatches[i])
            self.assertEqual(exp_hyperlinks[i],
                             hyperlink,
                             "Returned hyperlink and expected " +\
                             "hyperlink do not match!")
            self.assertEqual(exp_linknames[i],
                             link_name,
                             "Returned hyperlist and expected " +\
                             "hyperlist do not match!")

        # TESTING Hyperlinks with spaces!
        self.gr_api.errorlog = []
        regex_id = 0
        task_id = 2
        spec_name = 'Specname1'
        hypermatch = ('Storage of Data: ',
                      'file:',
                      '//\\en.tp.firm\\Diam$\\Com\\' +\
                      'SED_Dolore\\Magna-Aliquyam_PEDx' +\
                      '\x06%20lorem%20ipsum%20dolor%20sit' +\
                      '%20amet\x02%20labore%20magna\\P6\\' +\
                      'SED-Wizard\\SED-SPEC-Dummy',
                      '|Sed-Diam]]',
                      'Sed-Diam')

        hyperlink, link_name = \
            self.report.get_hyperlink(task_id, spec_name, regex_id, hypermatch)
        # IMPORTANT! use repr when printing
        #print"hyperlink:{}".format(repr(hyperlink))
        exp_hyperlink = 'file://\\en.tp.firm\\Diam$\\Com\\SED_Dolore' +\
            '\\Magna-Aliquyam_PEDx\x06%20lorem%20ipsum' +\
            '%20dolor%20sit%20amet\x02%20labore%20magna' +\
            '\\P6\\SED-Wizard\\SED-SPEC-Dummy'
        exp_link_name = "Sed-Diam"

        self.assertEqual(exp_hyperlink,
                         hyperlink,
                         "Returned hyperlink and expected " +\
                         "hyperlink do not match!")

        self.assertEqual(exp_link_name,
                         link_name,
                         "Returned hyperlist and expected " +\
                         "hyperlist do not match!")

        self.gr_api.errorlog = []
        regex_id = 0
        task_id = 3
        spec_name = 'Specname2'
        hypermatch = ('Storage of Data: ',
                      'file:',
                      '//\\en.tp.firm\\Diam$\\Com\\SED_Dolore' +\
                      '\\Magna-Aliquyam_PEDx\x06%20lorem%20' +\
                      'ipsum%20dolor%20sit%20amet\x02%20labore' +\
                      '%20magna\\P6\\SED-Wizard\\SED-SPEC-Dummy',
                      ']]',
                      '')
        hyperlink, link_name = \
            self.report.get_hyperlink(task_id, spec_name, regex_id, hypermatch)
        # IMPORTANT! use repr when printing
        #print"hyperlink:{}".format(repr(hyperlink))
        exp_hyperlink = 'file://\\en.tp.firm\\Diam$\\Com\\SED_Dolore' +\
            '\\Magna-Aliquyam_PEDx\x06%20lorem%20ipsum' +\
            '%20dolor%20sit%20amet\x02%20labore%20magna' +\
            '\\P6\\SED-Wizard\\SED-SPEC-Dummy'
        exp_link_name = 'file://\\en.tp.firm\\Diam$\\Com\\SED_Dolore' +\
            '\\Magna-Aliquyam_PEDx\x06%20lorem%20ipsum' +\
            '%20dolor%20sit%20amet\x02%20labore%20magna' +\
            '\\P6\\SED-Wizard\\SED-SPEC-Dummy'

        self.assertEqual(exp_hyperlink,
                         hyperlink,
                         "Returned hyperlink and expected " +\
                         "hyperlink do not match!")

        self.assertEqual(exp_link_name,
                         link_name,
                         "Returned hyperlist and expected " +\
                         "hyperlist do not match!")

    def test_get_hyperlink_ii(self):
        """ Test get_hyperlink errorlogs."""

        # TESTING errorlog for u'/wiki/'
        self.gr_api.errorlog = []
        regex_id = 1
        task_id = 2
        spec_name = 'Specname1'
        hypermatch = (u'', u'/wiki/', u'Specname1',
                      u' Specname1]', u'Specname1')
        self.report.get_hyperlink(task_id, spec_name, regex_id, hypermatch)
        errorlog = [("Specified link does not exist." +\
                     " Please check the full path and" +\
                     " the name of the following link:'Specname1']",
                     'http://example.org/Coconut/task/ticket/2',
                     'http://example.org/Coconut/event/wiki/Specname1')]

        self.assertEqual(self.gr_api.errorlog,
                         errorlog,
                         "Errorlogs do not match!")

        # TESTING errorlog for u'e:/wiki/'
        self.gr_api.errorlog = []
        regex_id = 1
        task_id = 2
        spec_name = 'Specname1'
        hypermatch = (u'', u'e:/wiki/', u'Specname1',
                      u' Specname1]', u'Specname1')
        self.report.get_hyperlink(task_id, spec_name, regex_id, hypermatch)
        errorlog = [("Specified link does not exist." +\
                     " Please check the full path and" +\
                     " the name of the following link:'Specname1']",
                     'http://example.org/Coconut/task/ticket/2',
                     'http://example.org/Coconut/event/wiki/Specname1')]

        self.assertEqual(self.gr_api.errorlog,
                         errorlog,
                         "Errorlogs do not match!")

        # TESTING errorlog for  u'wiki:'
        self.gr_api.errorlog = []
        regex_id = 1
        task_id = 2
        spec_name = 'Specname1'
        hypermatch = (u'', u'wiki:', u'/Specname1',
                      u' Specname1]', u'Specname1')
        self.report.get_hyperlink(task_id, spec_name, regex_id, hypermatch)
        errorlog = [("Specified link does not exist." +\
                     " Please check the full path and" +\
                     " the name of the following link:'Specname1']",
                     'http://example.org/Coconut/task/ticket/2',
                     'http://example.org/Coconut/event/wiki/Specname1')]

        self.assertEqual(self.gr_api.errorlog,
                         errorlog,
                         "Errorlogs do not match!")

    def test_get_hyperlink_iii(self):
        """ Test get_hyperlink NO errorlogs.
            There should be no errorlog
            because page exists."""

        pages = {
            'Specname1': """Lorem ipsum dolor sit amet,
consetetur sadipscing elitr, sed diam nonumy eirmod tempor
invidunt ut labore et dolore magna aliquyam erat, sed diam
voluptua."""}

        _insert_wiki_pages(self.gr_api.envs['event'], pages)

        # TESTING errorlog for u'/wiki/'
        self.gr_api.errorlog = []
        regex_id = 1
        task_id = 2
        spec_name = 'Specname1'
        hypermatch = (u'', u'/wiki/', u'Specname1',
                      u' Specname1]', u'Specname1')
        self.report.get_hyperlink(task_id, spec_name, regex_id, hypermatch)
        errorlog = []

        self.assertEqual(self.gr_api.errorlog,
                         errorlog,
                         "Errorlogs do not match!")

        # TESTING errorlog for u'e:/wiki/'
        self.gr_api.errorlog = []
        regex_id = 1
        task_id = 2
        spec_name = 'Specname1'
        hypermatch = (u'', u'e:/wiki/', u'Specname1',
                      u' Specname1]', u'Specname1')
        self.report.get_hyperlink(task_id, spec_name, regex_id, hypermatch)
        errorlog = []

        self.assertEqual(self.gr_api.errorlog,
                         errorlog,
                         "Errorlogs do not match!")

        # TESTING errorlog for  u'wiki:'
        self.gr_api.errorlog = []
        regex_id = 1
        task_id = 2
        spec_name = 'Specname1'
        hypermatch = (u'', u'wiki:', u'/Specname1',
                      u' Specname1]', u'Specname1')
        self.report.get_hyperlink(task_id, spec_name, regex_id, hypermatch)
        errorlog = []

        self.assertEqual(self.gr_api.errorlog,
                         errorlog,
                         "Errorlogs do not match!")

    def test_get_hyperlink_iv(self):
        """ Test get_hyperlink for examples:
        [[Dummy-APO-Database/GPD/Material_Strength| MS-GPD]],
        [[Dummy-APO-Database/GPD/Metallic_Joint| MBJ-GPD]]
        [[Dummy-APO-Database/GPD/Material_Strength]]."""

        # TESTING
        self.gr_api.errorlog = []
        regex_id = 4
        task_id = 2
        spec_name = 'APO/IP006/Dummy-APO-Database/IP006-APO-Spec-Sill'
        hypermatch = ('',
                      'IP006/',
                      'Dummy-APO-Database/GPD/Material_Strength',
                      '| MS-GPD]]',
                      ' MS-GPD')
        hyperlink, link_name = \
            self.report.get_hyperlink(task_id,
                                      spec_name,
                                      regex_id,
                                      hypermatch)

        exp_hyperlink = 'http://example.org/Coconut/event/' +\
            'wiki/APO/IP006/Dummy-APO-Database/GPD/Material_Strength'

        exp_link_name = ' MS-GPD'

        self.assertEqual(hyperlink,
                         exp_hyperlink,
                         "hyperlinks do not match!")

        self.assertEqual(link_name,
                         exp_link_name,
                         "linknames do not match!")

        # TESTING
        self.gr_api.errorlog = []
        regex_id = 4
        task_id = 2
        spec_name = 'APO/IP006/Dummy-APO-Database/IP006-APO-Spec-Sill'
        hypermatch = ('',
                      'IP006/',
                      'Dummy-APO-Database/GPD/Material_Strength',
                      ']]',
                      '')
        hyperlink, link_name = \
            self.report.get_hyperlink(task_id, spec_name, regex_id, hypermatch)

        exp_hyperlink = 'http://example.org/Coconut/event/' +\
            'wiki/APO/IP006/Dummy-APO-Database/GPD/Material_Strength'

        exp_link_name = 'IP006/Dummy-APO-Database/GPD/Material_Strength'

        self.assertEqual(hyperlink,
                         exp_hyperlink,
                         "hyperlinks do not match!")

        self.assertEqual(link_name,
                         exp_link_name,
                         "linknames do not match!")

    def test_get_merge_row(self): # pylint: disable=too-many-locals
        """ Test get_merge_row.
            params = [idr, row, table_row_length,
                      col_size, row_length, table,
                      task_id, spec_name]"""

        data = [[[' 1 '], [' 2 '], [' 3 ']],
                [[''], [' 1-2 '], [' 3 ']],
                [[' 1 '], [''], [' 2-3 ']],
                [['', ''], [' 1-2-3 ']]]

        merge_list = []
        table_row_length = set()

        col_size = 0

        task_id = 1

        spec_name = 'Specname 1'

        for item in data[0]:
            col_size += len(item)

        table = self.report.document.add_table(rows=1, cols=col_size)

        for idr, row in enumerate(data):
            row_length = 0

            for item in row:
                row_length += len(item)

            params = [idr,
                      row,
                      table_row_length,
                      col_size,
                      row_length,
                      table,
                      task_id,
                      spec_name]

            table, table_row_length, merge_row = \
                self.report.get_merge_row(params)

            merge_list.append(merge_row)
            merge_row = []
            row_length = 0

            if idr < len(data)-1:
                table.add_row()

        paragraph = self.report.get_paragraph_after_regex(r"Introduction")

        if paragraph is not None:
            new = paragraph.insert_paragraph_before()
            # _p and _tbl are protected and therefore not documented,
            # but otherwise it is impossible to insert a table,
            # and it would only be possible to append it to the
            # end of the document
            new._p.addnext(table._tbl) # pylint: disable=protected-access

        self.report.document.save(self.out)

        expected_table = [[[u' 1 ', u' 2 ', u' 3 '],
                           [u'', u' 1-2 ', u' 3 '],
                           [u' 1 ', u'', u' 2-3 '],
                           [u'', u'', u' 1-2-3 ']]]

        self.assertEqual(tables2txt(self.out),
                         expected_table,
                         "Returned table and expected " +\
                         "table do not match!")

        exp_merge_list = [[[]], [[0, 1]], [[1, 2]], [[0, 2]]]

        self.assertEqual(merge_list,
                         exp_merge_list,
                         "Returned merge list and expected " +\
                         "merge list do not match!")

    def test_find_merged_cells(self):
        """ Test find_merged_cells.
            args = [data,
                    table,
                    col_size,
                    task_id,
                    spec_name]"""

        data = [[[' 1 '], [' 2 '], [' 3 ']],
                [[''], [' 1-2 '], [' 3 ']],
                [[' 1 '], [''], [' 2-3 ']],
                [['', ''], [' 1-2-3 ']]]
        task_id = 1
        spec_name = 'Specname 1'
        merge_list = []
        col_size = 0

        for item in data[0]:
            col_size += len(item)

        table = self.report.document.add_table(rows=1, cols=col_size)

        args = [data,
                table,
                col_size,
                task_id,
                spec_name]

        table, merge_list = self.report.find_merged_cells(args)

        # As it can be seen from the data list above
        # No cell in the first row will be merged
        # Cells 0 and 1 in the second row will be merged
        # Cells 1 and 2 in the third row will be merged
        # Cells 0 to 2 in the fourth row will be merged

        expected_merge_list = [[[]], [[0, 1]], [[1, 2]], [[0, 2]]]

        self.assertEqual(merge_list,
                         expected_merge_list,
                         "Returned merge list and expected " +\
                         "merge list do not match!")

        paragraph = self.report.get_paragraph_after_regex(r"Introduction")

        if paragraph is not None:
            new = paragraph.insert_paragraph_before()
            # _p and _tbl are protected and therefore not documented,
            # but otherwise it is impossible to insert a table,
            # and it would only be possible to append it to the
            # end of the document
            new._p.addnext(table._tbl) # pylint: disable=protected-access

        self.report.document.save(self.out)

        expected_table = [[[u' 1 ', u' 2 ', u' 3 '],
                           [u'', u' 1-2 ', u' 3 '],
                           [u' 1 ', u'', u' 2-3 '],
                           [u'', u'', u' 1-2-3 ']]]

        self.assertEqual(tables2txt(self.out),
                         expected_table,
                         "Returned merge list and expected " +\
                         "merge list do not match!")

        self.gr_api.errorlog = []

        data = [[[' 1 '], [' 2 '], [' 3 ']],
                [[''], [' 1-2 '], [' 3 ']],
                [[' 1 '], [''], [' 2-3 ']],
                [['', '', ''], [' 1-2-3 ']]]

        merge_list = []
        col_size = 0

        for item in data[0]:
            col_size += len(item)

        table = self.report.document.add_table(rows=1, cols=col_size)

        args = [data,
                table,
                col_size,
                task_id,
                spec_name]

        table, merge_list = self.report.find_merged_cells(args)

        errorlog = [("There might be an extra pipe ||" +\
                     " in the wikitext of a table" +\
                     " that needs to be removed." +\
                     " Number of columns in each" +\
                     " row must match including merged" +\
                     " cells! Check the following" +\
                     " table with a: header: [[' 1 ']," +\
                     " [' 2 '], [' 3 ']]",
                     'http://example.org/Coconut/' +\
                     'task/ticket/1',
                     'http://example.org/Coconut/' +\
                     'event/wiki/Specname 1')]

        self.assertEqual(self.gr_api.errorlog,
                         errorlog,
                         "Errorlogs do not match!")

    def test_table_font_size(self): # pylint: disable=too-many-locals
        """ Test table_font_size."""

        data = [[[' 1 '], [' 2 '], [' 3 ']],
                [[''], [' 1-2 '], [' 3 ']],
                [[' 1 '], [''], [' 2-3 ']],
                [['', ''], [' 1-2-3 ']]]

        merge_list = [[[]],
                      [[0, 1]],
                      [[1, 2]],
                      [[0, 2]]]

        task_id = 1
        spec_name = 'Specname 1'

        col_size = 0

        for item in data[0]:
            col_size += len(item)

        table = self.report.document.add_table(rows=1, cols=col_size)

        args = [data,
                table,
                col_size,
                task_id,
                spec_name]

        table, merge_list = self.report.find_merged_cells(args)

        table = merge_table(table,
                            merge_list)

        table = table_font_size(table, 8)

        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font_size = run.font.size

        exp_font_size = 101600

        self.assertEqual(font_size,
                         exp_font_size,
                         "Table font sizes do not match")

        table = table_font_size(table, 14)

        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font_size = run.font.size

        exp_font_size = 177800

        self.assertEqual(font_size,
                         exp_font_size,
                         "Table font sizes do not match")

    def test_merge_table(self):
        """ Test merge_table."""

        data = [[[' 1 '], [' 2 '], [' 3 ']],
                [[''], [' 1-2 '], [' 3 ']],
                [[' 1 '], [''], [' 2-3 ']],
                [['', ''], [' 1-2-3 ']]]

        merge_list = [[[]],
                      [[0, 1]],
                      [[1, 2]],
                      [[0, 2]]]

        task_id = 1
        spec_name = 'Specname 1'

        col_size = 0

        for item in data[0]:
            col_size += len(item)

        table = self.report.document.add_table(rows=1, cols=col_size)

        args = [data,
                table,
                col_size,
                task_id,
                spec_name]

        table, merge_list = self.report.find_merged_cells(args)

        table = merge_table(table,
                            merge_list)

        paragraph = self.report.get_paragraph_after_regex(r"Introduction")

        if paragraph is not None:
            new = paragraph.insert_paragraph_before()
            # _p and _tbl are protected and therefore not documented,
            # but otherwise it is impossible to insert a table,
            # and it would only be possible to append it to the
            # end of the document
            new._p.addnext(table._tbl) # pylint: disable=protected-access

        self.report.document.save(self.out)

        expected_table = [[[u' 1 ', u' 2 ', u' 3 '],
                           [u' 1-2 ', u' 1-2 ', u' 3 '],
                           [u' 1 ', u' 2-3 ', u' 2-3 '],
                           [u' 1-2-3 ', u' 1-2-3 ', u' 1-2-3 ']]]

        self.assertEqual(tables2txt(self.out),
                         expected_table,
                         "Tables do not match in test_insert_table")

    def test_insert_table(self):
        """ Test insert_table and append_table."""

        data = [[[' 1 '], [' 2 '], [' 3 ']],
                [[''], [' 1-2 '], [' 3 ']],
                [[' 1 '], [''], [' 2-3 ']],
                [['', ''], [' 1-2-3 ']]]
        task_id = 1
        spec_name = 'Specname 1'

        paragraph = self.report.get_paragraph_after_regex(r"Introduction")

        self.report.insert_table(paragraph, data, task_id, spec_name)

        self.report.document.save(self.out)

        expected_table = [[[u' 1 ', u' 2 ', u' 3 '],
                           [u' 1-2 ', u' 1-2 ', u' 3 '],
                           [u' 1 ', u' 2-3 ', u' 2-3 '],
                           [u' 1-2-3 ', u' 1-2-3 ', u' 1-2-3 ']]]

        self.assertEqual(tables2txt(self.out),
                         expected_table,
                         "Tables do not match in test_insert_table")

    def test_get_table(self):
        """ Test get_table
            params = [i,
                      paragraph,
                      sections,
                      text,
                      spec_images] """

        sections = [
            [2,
             'Specname1',
             'consetetur sadipscing elitr,' +\
             'sed diam nonumy eirmod tempor\n' +\
             'http://www.test.com invidunt ut labore et dolore magna \n' +\
             '=== 1.2.1. Test Header 1 \n' +\
             '[/wiki/Specname1 Specname1] aliquyam erat, sed diam \n' +\
             'voluptua.\n',
             {},
             {'Table_11': [[[' 1 '], [' 2 '], [' 3 ']],
                           [[''], [' 1-2 '], [' 3 ']],
                           [[' 1 '], [''], [' 2-3 ']],
                           [['', ''], [' 1-2-3 ']]]}],
            [3,
             'Specname2',
             '[[Table(Table_21.tbl)]]\n' +\
             'duo dolores et ea rebum. Stet clita \n' +\
             '=== 1.2.2 Test Header 2 \n' +\
             'http://www.test.com kasd gubergren, no sea takimata \n' +\
             '[/wiki/Specname2 Specname2]' +\
             'sanctus est Lorem ipsum dolor sit amet.\n',
             {},
             {'Table_21': [[[' 1 '], [' 2 '], [' 3 ']],
                           [[''], [' 1-2 '], [' 3 ']],
                           [[' 1 '], [''], [' 2-3 ']],
                           [['', ''], [' 1-2-3 ']]]}]]

        paragraph = self.report.get_paragraph_after_regex(r"Introduction")

        level = 3

        spec_images = {}

        wiki_filter = \
            [re.compile(r'\s*\[\[Table\((.*)\.tbl\)\]\]\s*')]

        for i in range(len(sections)):
            text = sections[i][2]
            apo_spec = str(sections[i][0]) + ", " + sections[i][1]
            if self.report.report_type == 'SAR':
                apo_spec_path = sections[i][1].split('/')
                apo_spec = apo_spec_path[-1]
            style_key = 'Heading '+ str(level)
            paragraph.insert_paragraph_before(apo_spec, style=style_key)
            if text is not None:
                params = [i, paragraph, sections, text, spec_images]
                for line in params[3].splitlines():
                    line = to_unicode(line)
                    if wiki_filter[0].match(line):
                        self.report.get_table(\
                            params, wiki_filter[0].match(line).group(1))

        self.report.document.save(self.out)

        table = [[[u' 1 ', u' 2 ', u' 3 '],
                  [u' 1-2 ', u' 1-2 ', u' 3 '],
                  [u' 1 ', u' 2-3 ', u' 2-3 '],
                  [u' 1-2-3 ', u' 1-2-3 ', u' 1-2-3 ']]]

        self.assertEqual(tables2txt(self.out),
                         table,
                         "Tables do not match in " +\
                         "add_design_solutions of sar report!")

    def test_create_list(self):
        """ Test create_list."""

        text_list = ['**Step 2: Perform coordinate transformation for the' +\
                     ' side loads:** The provided input side load (GRA_Q1)',
                     ' Signs are reversed to match the Track2 GRA local' +\
                     ' Coordinate system (see [#Ref2] ([]))',
                     ' and then used for bolt load calculations. The sign' +\
                     ' convention used for the side loads “GRA (Q1)”' +\
                     ' is as follows:\\\\']

        task_id = 1

        spec_name = 'Specname 1'

        paragraph = \
            self.report.get_paragraph_after_regex(r"Sub-Component Description")

        for i in range(len(text_list)):
            paragraph = create_list(\
                paragraph.insert_paragraph_before(text=' ',
                                                  style='List Bullet'))
            line = filter_wiki_text(text_list[len(text_list)-1-i])
            line = to_unicode(line)
            args = [None,
                    paragraph,
                    line,
                    task_id,
                    spec_name]
            self.report.filter_hyperlinks(args)

        self.report.document.save(self.out)

        exp_doc = [u'APOS ANALYSED',
                   u'Introduction',
                   u'Sub-Component Description',
                   u' Step 2: Perform coordinate transformation' +\
                   ' for the side loads: The provided input' +\
                   ' side load (GRA_Q1)',
                   u'  Signs are reversed to match the Track2' +\
                   ' GRA local Coordinate system (see Ref2 ',
                   u'  and then used for bolt load calculations.' +\
                   u' The sign convention used for the side' +\
                   u' loads \u201cGRA (Q1)\u201d is as ',
                   u'Main Drawings Reference and Modification' +\
                   ' Applicability',
                   u'Material Properties',
                   u'Fastener Properties',
                   u'FEM Description (Global and Local)',
                   u'Load Cases and associated Criteria',
                   u'Main Hypotheses and Methods',
                   u'Applicable factors',
                   u'Documents',
                   u'Software']

        self.assertEqual(docx2txt(self.out),
                         exp_doc,
                         "Documents do not match in " +\
                         "test_create_list!")

    def test_find_sections(self):
        """ Test find_sections
            params = [i,
                      paragraph,
                      sections,
                      text,
                      spec_images] """

        sections = [
            [2,
             'Specname1',
             'consetetur sadipscing elitr,' +\
             'sed diam nonumy eirmod tempor\n' +\
             'http://www.test.com invidunt ut labore et dolore magna \n' +\
             '=== 1.2.1. Test Header 1 \n' +\
             '[/wiki/Specname1 Specname1] aliquyam erat, sed diam \n' +\
             'voluptua.\n',
             {},
             {'Table_11': [[[' 1 '], [' 2 '], [' 3 ']],
                           [[''], [' 1-2 '], [' 3 ']],
                           [[' 1 '], [''], [' 2-3 ']],
                           [['', ''], [' 1-2-3 ']]]}],
            [3,
             'Specname2',
             '[[Table(Table_21.tbl)]]\n' +\
             'duo dolores et ea rebum. Stet clita \n' +\
             '=== 1.2.2 Test Header 2 \n' +\
             'http://www.test.com kasd gubergren, no sea takimata \n' +\
             '[/wiki/Specname2 Specname2]' +\
             'sanctus est Lorem ipsum dolor sit amet.\n',
             {},
             {'Table_21': [[[' 1 '], [' 2 '], [' 3 ']],
                           [[''], [' 1-2 '], [' 3 ']],
                           [[' 1 '], [''], [' 2-3 ']],
                           [['', ''], [' 1-2-3 ']]]}]]

        paragraph = self.report.get_paragraph_after_regex(r"Introduction")

        level = 3

        spec_images = {}

        for i in range(len(sections)):
            text = sections[i][2]
            spec_images.update(sections[i][3])
            apo_spec = str(sections[i][0]) + ", " + sections[i][1]
            if self.report.report_type == 'SAR':
                apo_spec_path = sections[i][1].split('/')
                apo_spec = apo_spec_path[-1]
            style_key = 'Heading '+ str(level)
            paragraph.insert_paragraph_before(apo_spec, style=style_key)
            if text is not None:
                params = [i, paragraph, sections, text, spec_images]
                self.report.find_sections(params)

        self.report.document.save(self.out)

        doc = [u'APOS ANALYSED',
               u'Introduction',
               u'2, Specname1',
               u'consetetur sadipscing elitr,sed' +\
               ' diam nonumy eirmod tempor',
               u'invidunt ut labore et dolore magna ',
               u'Test Header 1',
               u' aliquyam erat, sed diam ',
               u'voluptua.',
               u'3, Specname2',
               u'',
               u'duo dolores et ea rebum. Stet clita ',
               u'Test Header 2',
               u'kasd gubergren, no sea takimata ',
               u'sanctus est Lorem ipsum dolor sit amet.',
               u'Sub-Component Description',
               u'Main Drawings Reference and' +\
               ' Modification Applicability',
               u'Material Properties',
               u'Fastener Properties',
               u'FEM Description (Global and Local)',
               u'Load Cases and associated Criteria',
               u'Main Hypotheses and Methods',
               u'Applicable factors',
               u'Documents',
               u'Software']

        self.assertEqual(docx2txt(self.out),
                         doc,
                         "Documents do not match in test_insert_section!")

        table = [[[u' 1 ', u' 2 ', u' 3 '],
                  [u' 1-2 ', u' 1-2 ', u' 3 '],
                  [u' 1 ', u' 2-3 ', u' 2-3 '],
                  [u' 1-2-3 ', u' 1-2-3 ', u' 1-2-3 ']]]

        self.assertEqual(tables2txt(self.out),
                         table,
                         "Tables do not match in " +\
                         "add_design_solutions of sar report!")

    def test_insert_section(self):
        """ Test insert_section and filter_spec_hyperlinks."""

        pages = {
            'Specname1': """consetetur sadipscing elitr,
sed diam nonumy eirmod tempor
http://www.test.com invidunt ut labore et dolore magna 
[/wiki/Specname1 Specname1] aliquyam erat, sed diam 
voluptua.""",
            'Specname2': """[[Table(Table_21.tbl)]]
duo dolores et ea rebum. Stet clita 
http://www.test.com kasd gubergren, no sea takimata 
[/wiki/Specname2 Specname2]
sanctus est Lorem ipsum dolor sit amet."""}

        _insert_wiki_pages(self.gr_api.envs['event'], pages)

        sections = [
            [2,
             'Specname1',
             'consetetur sadipscing elitr,' +\
             'sed diam nonumy eirmod tempor\n' +\
             'http://www.test.com invidunt ut labore et dolore magna \n' +\
             '=== 1.2.1. Test Header 1 \n' +\
             '[/wiki/Specname1 Specname1] aliquyam erat, sed diam \n' +\
             'voluptua.\n',
             {},
             {'Table_11': [[[' 1 '], [' 2 '], [' 3 ']],
                           [[''], [' 1-2 '], [' 3 ']],
                           [[' 1 '], [''], [' 2-3 ']],
                           [['', ''], [' 1-2-3 ']]]}],
            [3,
             'Specname2',
             '[[Table(Table_21.tbl)]]\n' +\
             'duo dolores et ea rebum. Stet clita \n' +\
             '=== 1.2.2 Test Header 2 \n' +\
             'http://www.test.com kasd gubergren, no sea takimata \n' +\
             '[/wiki/Specname2 Specname2]' +\
             'sanctus est Lorem ipsum dolor sit amet.\n',
             {},
             {'Table_21': [[[' 1 '], [' 2 '], [' 3 ']],
                           [[''], [' 1-2 '], [' 3 ']],
                           [[' 1 '], [''], [' 2-3 ']],
                           [['', ''], [' 1-2-3 ']]]}]]

        paragraph = self.report.get_paragraph_after_regex(r"Introduction")

        self.report.insert_section(paragraph, sections, 3)

        self.report.document.save(self.out)

        doc = [u'APOS ANALYSED',
               u'Introduction',
               u'2, Specname1',
               u'consetetur sadipscing elitr' +\
               ',sed diam nonumy eirmod tempor',
               u'invidunt ut labore et dolore magna ',
               u'Test Header 1',
               u' aliquyam erat, sed diam ',
               u'voluptua.',
               u'3, Specname2',
               u'',
               u'duo dolores et ea rebum. Stet clita ',
               u'Test Header 2',
               u'kasd gubergren, no sea takimata ',
               u'sanctus est Lorem ipsum dolor sit amet.',
               u'Sub-Component Description',
               u'Main Drawings Reference and' +\
               ' Modification Applicability',
               u'Material Properties',
               u'Fastener Properties',
               u'FEM Description (Global and Local)',
               u'Load Cases and associated Criteria',
               u'Main Hypotheses and Methods',
               u'Applicable factors',
               u'Documents',
               u'Software']

        self.assertEqual(docx2txt(self.out),
                         doc,
                         "Documents do not match in test_insert_section!")

        table = [[[u' 1 ', u' 2 ', u' 3 '],
                  [u' 1-2 ', u' 1-2 ', u' 3 '],
                  [u' 1 ', u' 2-3 ', u' 2-3 '],
                  [u' 1-2-3 ', u' 1-2-3 ', u' 1-2-3 ']]]

        self.assertEqual(tables2txt(self.out),
                         table,
                         "Tables do not match in " +\
                         "add_design_solutions of sar report!")

    def test_add_introduction(self):
        """ Test add_introduction, note that if out_adc.docx
            and out_sar.docx files are not deleted, it can be
            observed that links are printed in out_adc.docx but
            not in out_sar.docx."""

        report_adc_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'2, Specname1',
                              u'',
                              u'consetetur sadipscing elitr' +\
                              ', sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'3, Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Sub-Component Description',
                              u'Main Drawings Reference and' +\
                              ' Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Storage of Data']

        report_sar_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Specname1',
                              u'',
                              u'consetetur sadipscing elitr' +\
                              ', sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Sub-Component Description',
                              u'Main Drawings Reference and' +\
                              ' Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Abbreviations and Units']

        self.report_adc.add_introduction(self.section)
        self.report_sar.add_introduction(self.section)

        self.report_adc.save(self.out_adc)
        self.report_sar.save(self.out_sar)

        self.assertEqual(docx2txt(self.out_adc),
                         report_adc_section,
                         "Sections do not match in " +\
                         "add_introduction of adc report!")
        self.assertEqual(tables2txt(self.out_adc),
                         self.tables,
                         "Tables do not match in " +\
                         "add_introduction of adc report!")

        self.assertEqual(docx2txt(self.out_sar),
                         report_sar_section,
                         "Sections do not match in " +\
                         "add_introduction of sar report!")
        self.assertEqual(tables2txt(self.out_sar),
                         self.tables,
                         "Tables do not match in " +\
                         "add_introduction of sar report!")

    def test_add_structural_function(self):
        """ Test add_structural_function, note that if out_adc.docx
            and out_sar.docx files are not deleted, it can be
            observed that links are printed in out_adc.docx but
            not in out_sar.docx."""

        report_adc_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'2, Specname1',
                              u'',
                              u'consetetur sadipscing elitr' +\
                              ', sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'3, Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Main Drawings Reference and' +\
                              ' Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Storage of Data']

        report_sar_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Specname1',
                              u'',
                              u'consetetur sadipscing elitr' +\
                              ', sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Main Drawings Reference and' +\
                              ' Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Abbreviations and Units']

        self.report_adc.add_structural_function(self.section)
        self.report_sar.add_structural_function(self.section)

        self.report_adc.save(self.out_adc)
        self.report_sar.save(self.out_sar)

        self.assertEqual(docx2txt(self.out_adc),
                         report_adc_section,
                         "Sections do not match in " +\
                         "add_structural_function of adc report!")
        self.assertEqual(tables2txt(self.out_adc),
                         self.tables,
                         "Tables do not match in " +\
                         "add_structural_function of adc report!")

        self.assertEqual(docx2txt(self.out_sar),
                         report_sar_section,
                         "Sections do not match in " +\
                         "add_structural_function of sar report!")
        self.assertEqual(tables2txt(self.out_sar),
                         self.tables,
                         "Tables do not match in " +\
                         "add_structural_function of sar report!")

    def test_add_design_solutions(self):
        """ Test add_design_solutions, note that if out_adc.docx
            and out_sar.docx files are not deleted, it can be
            observed that links are printed in out_adc.docx but
            not in out_sar.docx."""

        report_adc_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference and' +\
                              ' Modification Applicability',
                              u'2, Specname1',
                              u'',
                              u'consetetur sadipscing elitr' +\
                              ', sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'3, Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Storage of Data']

        report_sar_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference' +\
                              ' and Modification Applicability',
                              u'Specname1',
                              u'',
                              u'consetetur sadipscing elitr' +\
                              ', sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Abbreviations and Units']

        self.report_adc.add_design_solutions(self.section)
        self.report_sar.add_design_solutions(self.section)

        self.report_adc.save(self.out_adc)
        self.report_sar.save(self.out_sar)

        self.assertEqual(docx2txt(self.out_adc),
                         report_adc_section,
                         "Sections do not match in " +\
                         "add_design_solutions of adc report!")
        self.assertEqual(tables2txt(self.out_adc),
                         self.tables,
                         "Tables do not match in " +\
                         "add_design_solutions of adc report!")

        self.assertEqual(docx2txt(self.out_sar),
                         report_sar_section,
                         "Sections do not match in " +\
                         "add_design_solutions of sar report!")
        self.assertEqual(tables2txt(self.out_sar),
                         self.tables,
                         "Tables do not match in " +\
                         "add_design_solutions of sar report!")

    def test_add_material_data(self):
        """ Test add_material_data, note that if out_adc.docx
            and out_sar.docx files are not deleted, it can be
            observed that links are printed in out_adc.docx but
            not in out_sar.docx."""

        report_adc_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference' +\
                              ' and Modification Applicability',
                              u'Material Properties',
                              u'2, Specname1',
                              u'',
                              u'consetetur sadipscing elitr' +\
                              ', sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'3, Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Storage of Data']

        report_sar_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference' +\
                              ' and Modification Applicability',
                              u'Material Properties',
                              u'Specname1',
                              u'',
                              u'consetetur sadipscing elitr' +\
                              ', sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Abbreviations and Units']

        self.report_adc.add_material_data(self.section)
        self.report_sar.add_material_data(self.section)

        self.report_adc.save(self.out_adc)
        self.report_sar.save(self.out_sar)

        self.assertEqual(docx2txt(self.out_adc),
                         report_adc_section,
                         "Sections do not match in " +\
                         "add_material_data of adc report!")
        self.assertEqual(tables2txt(self.out_adc),
                         self.tables,
                         "Tables do not match in " +\
                         "add_material_data of adc report!")

        self.assertEqual(docx2txt(self.out_sar),
                         report_sar_section,
                         "Sections do not match in " +\
                         "add_material_data of sar report!")
        self.assertEqual(tables2txt(self.out_sar),
                         self.tables,
                         "Tables do not match in " +\
                         "add_material_data of sar report!")

    def test_add_fastener_data(self):
        """ Test add_fastener_data, note that if out_adc.docx
            and out_sar.docx files are not deleted, it can be
            observed that links are printed in out_adc.docx but
            not in out_sar.docx."""

        report_adc_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference and' +\
                              ' Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'2, Specname1',
                              u'',
                              u'consetetur sadipscing elitr,' +\
                              ' sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'3, Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Storage of Data']

        report_sar_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference and' +\
                              ' Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'Specname1',
                              u'',
                              u'consetetur sadipscing elitr,' +\
                              ' sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Abbreviations and Units']

        self.report_adc.add_fastener_data(self.section)
        self.report_sar.add_fastener_data(self.section)

        self.report_adc.save(self.out_adc)
        self.report_sar.save(self.out_sar)

        self.assertEqual(docx2txt(self.out_adc),
                         report_adc_section,
                         "Sections do not match in " +\
                         "add_fastener_data of adc report!")
        self.assertEqual(tables2txt(self.out_adc),
                         self.tables,
                         "Tables do not match in " +\
                         "add_fastener_data of adc report!")

        self.assertEqual(docx2txt(self.out_sar),
                         report_sar_section,
                         "Sections do not match in " +\
                         "add_fastener_data of sar report!")
        self.assertEqual(tables2txt(self.out_sar),
                         self.tables,
                         "Tables do not match in " +\
                         "add_fastener_data of sar report!")

    def test_add_applicable_fems(self):
        """ Test add_applicable_fems, note that if out_adc.docx
            and out_sar.docx files are not deleted, it can be
            observed that links are printed in out_adc.docx but
            not in out_sar.docx."""

        report_adc_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference and' +\
                              ' Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'2, Specname1',
                              u'',
                              u'consetetur sadipscing elitr,' +\
                              ' sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'3, Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Storage of Data']

        report_sar_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference and' +\
                              ' Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Specname1',
                              u'',
                              u'consetetur sadipscing elitr,' +\
                              ' sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Abbreviations and Units']

        self.report_adc.add_applicable_fems(self.section)
        self.report_sar.add_applicable_fems(self.section)

        self.report_adc.save(self.out_adc)
        self.report_sar.save(self.out_sar)

        self.assertEqual(docx2txt(self.out_adc),
                         report_adc_section,
                         "Sections do not match in " +\
                         "add_applicable_fems of adc report!")
        self.assertEqual(tables2txt(self.out_adc),
                         self.tables,
                         "Tables do not match in " +\
                         "add_applicable_fems of adc report!")

        self.assertEqual(docx2txt(self.out_sar),
                         report_sar_section,
                         "Sections do not match in " +\
                         "add_applicable_fems of sar report!")
        self.assertEqual(tables2txt(self.out_sar),
                         self.tables,
                         "Tables do not match in " +\
                         "add_applicable_fems of sar report!")

    def test_add_applicable_load_cases(self):
        """ Test add_applicable_load_cases, note that if out_adc.docx
            and out_sar.docx files are not deleted, it can be
            observed that links are printed in out_adc.docx but
            not in out_sar.docx."""

        report_adc_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference and' +\
                              ' Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'2, Specname1',
                              u'',
                              u'consetetur sadipscing elitr,' +\
                              ' sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'3, Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Storage of Data']

        report_sar_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference and'+\
                              ' Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Specname1',
                              u'',
                              u'consetetur sadipscing elitr, sed' +\
                              ' diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Abbreviations and Units']

        self.report_adc.add_applicable_load_cases(self.section)
        self.report_sar.add_applicable_load_cases(self.section)

        self.report_adc.save(self.out_adc)
        self.report_sar.save(self.out_sar)

        self.assertEqual(docx2txt(self.out_adc),
                         report_adc_section,
                         "Sections do not match in " +\
                         "add_applicable_load_cases of adc report!")
        self.assertEqual(tables2txt(self.out_adc),
                         self.tables,
                         "Tables do not match in " +\
                         "add_applicable_load_cases of adc report!")

        self.assertEqual(docx2txt(self.out_sar),
                         report_sar_section,
                         "Sections do not match in " +\
                         "add_applicable_load_cases of sar report!")
        self.assertEqual(tables2txt(self.out_sar),
                         self.tables,
                         "Tables do not match in " +\
                         "add_applicable_load_cases of sar report!")

    def test_add_sizing_criteria(self):
        """ Test add_sizing_criteria, note that if out_adc.docx
            and out_sar.docx files are not deleted, it can be
            observed that links are printed in out_adc.docx but
            not in out_sar.docx."""

        report_adc_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference and' +\
                              ' Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'2, Specname1',
                              u'',
                              u'consetetur sadipscing elitr,' +\
                              ' sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'3, Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Storage of Data']

        report_sar_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference and' +\
                              ' Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Specname1',
                              u'',
                              u'consetetur sadipscing elitr,' +\
                              ' sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Abbreviations and Units']

        self.report_adc.add_sizing_criteria(self.section)
        self.report_sar.add_sizing_criteria(self.section)

        self.report_adc.save(self.out_adc)
        self.report_sar.save(self.out_sar)

        self.assertEqual(docx2txt(self.out_adc),
                         report_adc_section,
                         "Sections do not match in " +\
                         "add_sizing_criteria of adc report!")
        self.assertEqual(tables2txt(self.out_adc),
                         self.tables,
                         "Tables do not match in " +\
                         "add_sizing_criteria of adc report!")

        self.assertEqual(docx2txt(self.out_sar),
                         report_sar_section,
                         "Sections do not match in " +\
                         "add_applicable_load_cases of sar report!")
        self.assertEqual(tables2txt(self.out_sar),
                         self.tables,
                         "Tables do not match in " +\
                         "add_applicable_load_cases of sar report!")

    def test_add_applicable_factors(self):
        """ Test add_applicable_factors, note that if out_adc.docx
            and out_sar.docx files are not deleted, it can be
            observed that links are printed in out_adc.docx but
            not in out_sar.docx."""

        report_adc_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference and' +\
                              ' Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'2, Specname1',
                              u'',
                              u'consetetur sadipscing elitr,' +\
                              ' sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'3, Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Documents',
                              u'Software',
                              u'Storage of Data']

        report_sar_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference and' +\
                              ' Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Specname1',
                              u'',
                              u'consetetur sadipscing elitr,' +\
                              ' sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Documents',
                              u'Software',
                              u'Abbreviations and Units']

        self.report_adc.add_applicable_factors(self.section)
        self.report_sar.add_applicable_factors(self.section)

        self.report_adc.save(self.out_adc)
        self.report_sar.save(self.out_sar)

        self.assertEqual(docx2txt(self.out_adc),
                         report_adc_section,
                         "Sections do not match in " +\
                         "add_applicable_factors of adc report!")
        self.assertEqual(tables2txt(self.out_adc),
                         self.tables,
                         "Tables do not match in " +\
                         "add_applicable_factors of adc report!")

        self.assertEqual(docx2txt(self.out_sar),
                         report_sar_section,
                         "Sections do not match in " +\
                         "add_applicable_factors of sar report!")
        self.assertEqual(tables2txt(self.out_sar),
                         self.tables,
                         "Tables do not match in " +\
                         "add_applicable_factors of sar report!")

    def test_add_documents(self):
        """ Test add_documents, note that if out_adc.docx
            and out_sar.docx files are not deleted, it can be
            observed that links are printed in out_adc.docx but
            not in out_sar.docx."""

        report_adc_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference' +\
                              ' and Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'2, Specname1',
                              u'',
                              u'consetetur sadipscing elitr,' +\
                              ' sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'3, Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Software',
                              u'Storage of Data']

        report_sar_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference' +\
                              ' and Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Specname1',
                              u'',
                              u'consetetur sadipscing elitr,' +\
                              ' sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Software',
                              u'Abbreviations and Units']

        self.report_adc.add_documents(self.section)
        self.report_sar.add_documents(self.section)

        self.report_adc.save(self.out_adc)
        self.report_sar.save(self.out_sar)

        self.assertEqual(docx2txt(self.out_adc),
                         report_adc_section,
                         "Sections do not match in " +\
                         "add_documents of adc report!")
        self.assertEqual(tables2txt(self.out_adc),
                         self.tables,
                         "Tables do not match in " +\
                         "add_documents of adc report!")

        self.assertEqual(docx2txt(self.out_sar),
                         report_sar_section,
                         "Sections do not match in " +\
                         "add_documents of sar report!")
        self.assertEqual(tables2txt(self.out_sar),
                         self.tables,
                         "Tables do not match in " +\
                         "add_documents of sar report!")

    def test_add_software(self):
        """ Test add_software, note that if out_adc.docx
            and out_sar.docx files are not deleted, it can be
            observed that links are printed in out_adc.docx but
            not in out_sar.docx."""

        report_adc_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference' +\
                              ' and Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'2, Specname1',
                              u'',
                              u'consetetur sadipscing elitr' +\
                              ', sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'3, Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Storage of Data']

        report_sar_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference' +\
                              ' and Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Specname1',
                              u'',
                              u'consetetur sadipscing elitr' +\
                              ', sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Abbreviations and Units']

        self.report_adc.add_software(self.section)
        self.report_sar.add_software(self.section)

        self.report_adc.save(self.out_adc)
        self.report_sar.save(self.out_sar)

        self.assertEqual(docx2txt(self.out_adc),
                         report_adc_section,
                         "Sections do not match in " +\
                         "add_software of adc report!")
        self.assertEqual(tables2txt(self.out_adc),
                         self.tables,
                         "Tables do not match in " +\
                         "add_software of adc report!")

        self.assertEqual(docx2txt(self.out_sar),
                         report_sar_section,
                         "Sections do not match in " +\
                         "add_software of sar report!")
        self.assertEqual(tables2txt(self.out_sar),
                         self.tables,
                         "Tables do not match in " +\
                         "add_software of sar report!")

    def test_insert_analysed_apos_table(self):
        """ Test insert_analysed_apos_table, note that if out_adc.docx
            and out_sar.docx files are not deleted, it can be
            observed that links are printed in out_adc.docx but
            not in out_sar.docx."""

        sections = [[2, u'Summary 1',
                     'Storage of Data: [[http://localhost/' +\
                     'Coconut/event/wiki/Specname1|Link1]]'],
                    [3, u'Summary 2',
                     'Storage of Data: [[http://localhost/' +\
                     'Coconut/event/wiki/Specname2|Link2]]'],
                    [4, u'Summary 3',
                     'Storage of Data: [[http://localhost/' +\
                     'Coconut/event/wiki/Specname3|Link3]]'],
                    [5, u'Summary 4',
                     'Storage of Data: [[http://localhost/' +\
                     'Coconut/event/wiki/Specname4|Link4]]']]

        tables = [[[u'ITEM', u'ANALYSE APO TASK No',
                    u'APO TASK NAME', u'REMARK'],
                   [u'1', u'2', u'Summary 1', u''],
                   [u'2', u'3', u'Summary 2', u''],
                   [u'3', u'4', u'Summary 3', u''],
                   [u'4', u'5', u'Summary 4', u'']]]

        report_adc_section = [u'APOS ANALYSED',
                              u'',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference' +\
                              ' and Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Storage of Data']

        self.report_adc.add_analysed_apos(sections)
        self.report_adc.save(self.out_adc)

        self.assertEqual(docx2txt(self.out_adc),
                         report_adc_section,
                         "Sections do not match in " +\
                         "add_storage_of_data of adc report!")
        self.assertEqual(tables2txt(self.out_adc),
                         tables,
                         "Tables do not match in " +\
                         "add_storage_of_data of adc report!")

    def test_add_storage_of_data(self):
        """ Test add_storage_of_data, note that if out_adc.docx
            and out_sar.docx files are not deleted, it can be
            observed that links are printed in out_adc.docx but
            not in out_sar.docx."""

        report_adc_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference' +\
                              ' and Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Storage of Data',
                              u'2, Specname1',
                              u'',
                              u'consetetur sadipscing elitr' +\
                              ', sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'3, Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'']

        self.report_adc.add_storage_of_data(self.section)
        self.report_adc.save(self.out_adc)

        self.assertEqual(docx2txt(self.out_adc),
                         report_adc_section,
                         "Sections do not match in " +\
                         "add_storage_of_data of adc report!")
        self.assertEqual(tables2txt(self.out_adc),
                         self.tables,
                         "Tables do not match in " +\
                         "add_storage_of_data of adc report!")

    def test_add_abbreviations(self):
        """ Test add_abbreviations, note that if out_adc.docx
            and out_sar.docx files are not deleted, it can be
            observed that links are printed in out_adc.docx but
            not in out_sar.docx."""

        report_sar_section = [u'APOS ANALYSED',
                              u'Introduction',
                              u'Sub-Component Description',
                              u'Main Drawings Reference' +\
                              ' and Modification Applicability',
                              u'Material Properties',
                              u'Fastener Properties',
                              u'FEM Description (Global and Local)',
                              u'Load Cases and associated Criteria',
                              u'Main Hypotheses and Methods',
                              u'Applicable factors',
                              u'Documents',
                              u'Software',
                              u'Specname1',
                              u'',
                              u'consetetur sadipscing elitr' +\
                              ', sed diam nonumy eirmod tempor',
                              u'invidunt ut labore et dolore magna ',
                              u' aliquyam erat, sed diam ',
                              u'voluptua.',
                              u'Specname2',
                              u'',
                              u'duo dolores et ea rebum. Stet clita ',
                              u'() kasd gubergren, no sea takimata ',
                              u'sanctus est Lorem ipsum dolor sit amet.',
                              u'Abbreviations and Units']

        self.report_sar.add_software(self.section)
        self.report_sar.save(self.out_sar)

        self.assertEqual(docx2txt(self.out_sar),
                         report_sar_section,
                         "Sections do not match in " +\
                         "add_abbreviations of sar report!")
        self.assertEqual(tables2txt(self.out_sar),
                         self.tables,
                         "Tables do not match in " +\
                         "add_abbreviations of sar report!")

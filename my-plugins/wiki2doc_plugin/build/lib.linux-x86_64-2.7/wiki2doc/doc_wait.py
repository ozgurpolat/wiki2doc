""" class to create a word document """

import re
import os
import tempfile
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT

from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
#from trac.mimeview import Context
from trac.web.chrome import web_context
from trac.util.text import to_unicode
from trac.wiki.formatter import HtmlFormatter
from HTMLParser import HTMLParser
from helpers import\
add_hyperlink,\
check_table_row_length,\
check_string,\
create_list,\
filter_wiki_text,\
find_hyperlinks,\
get_base_url,\
get_link_name,\
get_wiki_specname,\
insert_image,\
merge_table,\
process_blockquote,\
remove_forward_slash,\
table_font_size,\
get_header_in_text_line

from parser import DocumentHTMLParser

class Doc(object): # pylint: disable=too-many-public-methods
    """ class to create a document in MS Word """

    def __init__(self, args):
        self.document = Document(args[0])
        self.env = args[1]
        self.wiki2doc = args[2]
        self.req = args[3]
        
        self.add_hyper_link = True

    def insert_paragraph_after(self, paragraph, text=None, style=None):
        """Insert a new paragraph after the given paragraph."""
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        if style is not None:
            new_para.style = style
        return new_para

    def save(self, path):
        """ save docx to path """
        self.document.save(path)

    def get_content(self):
         """ save docx to path """
         _, out = tempfile.mkstemp()
         self.save(out)
         with open(out) as filehndl:
             content = filehndl.read()
         os.unlink(out)
         return content
  
    def get_paragraph_after_regex(self, regex):
         """ helper function to be used before insert_paragraph_before() """
  
         regex = re.compile(regex)
         found = False
         idx = 0
         while not found and idx < len(self.document.paragraphs):
             par = self.document.paragraphs[idx]
             match = regex.match(par.text)
             if match:
                 found = True
             else:
                 idx += 1
         idx += 1
         if found:
             if idx < len(self.document.paragraphs):
                 return self.document.paragraphs[idx]
         return self.document.add_paragraph()
  
    def get_merge_row(self, params):
         """ for a given table data, analyses
             the data and finds merged cells.
             params = [idr, row, table_row_length,
                       col_size, row_length, table,
                       spec_name]"""
  
         merge_row = []
         col = 0
         pos = 0
         start = 0
         end = 0
  
         print('inside get_merge_row params:', params)  
  
         for idx, item in enumerate(params[1]):
             for idy, value in enumerate(item):
                 if check_table_row_length(params[3],
                                           params[4]):
                     value = filter_wiki_text(value)
                     #args = [table, paragraph,
                     #        text, task_id, spec_name]
                     args = [params[5],
                             params[5].rows[params[0]].cells[col].paragraphs[0],
                             value,
                             params[6]]
                     #args[2].rows[args[0]].cells[args[1]].paragraphs[0]
                     params[5], _ = self.filter_hyperlinks(args)
                 else:
                     params[2].add(False)
                 col += 1
                 start = pos
                 end = pos + len(item)
                 if idy == 0 and value == '' and idx < len(params[1])-1:
                     merge_row.append([start, end])
                 elif idy == 0 and value == '' and \
                     idx == len(params[1])-1 and len(item) > 1:
                     merge_row.append([start, end-1])
                 elif idy == 0 and col == params[4] - 1 and \
                     len(merge_row) == 0:
                     merge_row.append([])
             pos += len(item)
  
         return (params[5], params[2], merge_row)
  
  
    def find_merged_cells(self, args):
        """ for a given table data, analyses
            the data and finds merged cells.
            args = [data,
                    table,
                    col_size,
                    spec_name]"""
        
        print('inside find_merged_cells args:', args)            
        
        merge_list = []
        table_row_length = set()
        for idr, row in enumerate(args[0]):
            row_length = 0
  
            for item in row:
                row_length += len(item)
            params = [idr,
                      row,
                      table_row_length,
                      args[2],
                      row_length,
                      args[1],
                      args[3]]
            
            print('inside find_merged_cells params:', params)     
            
            args[1], table_row_length, merge_row = self.get_merge_row(params)
  
            merge_list.append(merge_row)
            merge_row = []
            row_length = 0
  
            if idr < len(args[0])-1:
                args[1].add_row()
  
        if len(list(table_row_length)) > 0:
            page_path = get_base_url(self.req) +\
                'wiki/' + args[3]
            self.wiki2doc.errorlog.append((
                "There might be an extra pipe || in the wikitext of" +\
                " a table that needs to be removed. Number of" +\
                " columns in each row must match including merged" +\
                " cells! Check the following table with a:" +\
                " header: {}".format(args[0][0]),
                page_path))
        return (args[1], merge_list)
  
    def append_table(self, data, spec_name):
        """ for a given table data, this function analyzes the table,
            looks for cells to be merged inside the text as described
            below. Creates the table with the values first, then calls
            merge_table method to merge the cells if they need to be merged,
            then returns the table.
  
            Wiki Markup:
            || 1 || 2 || 3 ||
            |||| 1-2 || 3 ||
            || 1 |||| 2-3 ||
            |||||| 1-2-3 ||
  
            Display:
            ---- --- ----
            | 1 | 2 | 3 |
            ---- --- ----
            | 1-2   | 3 |
            -------- ----
            | 1 | 2-3   |
            -------------
            | 1-2-3     |
            -------------"""
        merge_list = []
        col_size = 0
  
        for item in data[0]:
            col_size += len(item)
  
        table = self.document.add_table(rows=1, cols=col_size)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER # pylint: disable=no-member
        args = [data,
                table,
                col_size,
                spec_name]
        table, merge_list = self.find_merged_cells(args)
  
        if len(data) == len(merge_list):
            table = merge_table(table, merge_list)
            table = table_font_size(table, 8)
            return table
        else:
            self.wiki2doc.errorlog.append(
                "Merge cell list length and table length does not match." +\
                "Please check the merged cells in: {} \n".format(data[0]),
                'None')
  
    def insert_table(self, paragraph, table_data, spec_name):
        """ insert table """
  
        # ************************************* IMPORTANT *******
        # ADD THIS! KeyError: u"no style with name 'TableGrid'"
        '''
        Trac detected an internal error:
        KeyError: u"no style with name TableGrid
        There was an internal error in Trac. It is recommended that you notify your local Trac administrator with the information needed to reproduce the issue.
        
        To that end, you could  a ticket.
        
        The action that triggered the error was:
        
        POST: /wiki2doc
        TracGuide The Trac User and Administration Guide        
        '''
  
        table = self.append_table(table_data, spec_name)
        table.style = 'Table Grid'
  
        if paragraph is not None:
            new = paragraph.insert_paragraph_before()
            # _p and _tbl are protected and therefore not documented,
            # but otherwise it is impossible to insert a table,
            # and it would only be possible to append it to the
            # end of the document
            new._p.addnext(table._tbl) # pylint: disable=protected-access
  
    def errorlog_missing_page(self, missing_spec, in_spec):
        """ Errorlog a wiki page that
            does not exist. """
        missing_spec = remove_forward_slash(missing_spec)
        self.wiki2doc.errorlog.append(\
            ("Specified link does not exist. Please check the " +\
             "full path and the name of the following link:'{}']".format(\
                                                           missing_spec),
             get_base_url(self.req) + 'wiki/' + in_spec))
  
    def get_wiki_hyperlink(self, spec_name, hyper):
        """ returns the wiki page hyperlink path for
            another page that is under same parent
            path. See regex_id 4. in find_hyperlinks.
  
            Example wikipage: http://10.45.43.145:8000/Coconut/
            event/wiki/APO/IP006/Dummy-APO-Database/IP006-APO-Spec-Sill
  
            Example reference from inside the link above.
            [[Dummy-APO-Database/GPD/Material_Strength| MS-GPD]]
            [[Dummy-APO-Database/GPD/Metallic_Joint| MBJ-GPD]]
            [[/GPD/Material_Strength| MS-GPD]]
            [[GPD/Material_Strength| MS-GPD]]
            [[/IP006/Dummy-APO-Database/GPD/Material_Strength| MS-GPD]]
            [[IP006/Dummy-APO-Database/GPD/Material_Strength| MS-GPD]]
            [[IP006/Dummy-APO-Database/GPD/Material_Strength]]
  
            This works because both "IP006-APO-Spec-Sill" and
            "GPD/Material_Strength" are under:
  
            http://10.45.43.145:8000/Coconut/
            event/wiki/APO/IP006/Dummy-APO-Database/
            """
  
        given_path = remove_forward_slash(hyper[1]) + hyper[2]
        given_path_list = given_path.split("/")
        full_wiki_path = get_base_url(self.req) +\
            "wiki/" + spec_name
        full_path_list = full_wiki_path.split("/")
        protocol = full_path_list[0]
        full_path_list = full_path_list[2:]
        list_index = []
        hyperlink = ''
  
        for i, item in enumerate(full_path_list):
            if item in set(given_path_list):
                list_index.append(i)
  
        if len(list_index) > 0:
            full_path_list = full_path_list[:list_index[0]]
        elif len(list_index) == 0:
            full_path_list = full_path_list[:-1]
  
        mod_full_path = ''
  
        for item in full_path_list:
            mod_full_path += item + "/"
  
        hyperlink = protocol + "//" + mod_full_path + given_path
  
        return hyperlink
  
    def get_link_path(self, spec_name, regex_id, hyper):
        """ for a given hypermatch this function
            returns the hyperlink path."""
  
        hyperlink = ''
  
        if regex_id == 4:
            another_child_spec_name = get_wiki_specname(spec_name, hyper)
            page = \
                self.wiki2doc.get_wikipage(
                    remove_forward_slash(another_child_spec_name))
            if not page:
                self.errorlog_missing_page(another_child_spec_name,
                                           spec_name)
            hyperlink = self.get_wiki_hyperlink(spec_name, hyper)
        elif hyper[1] == '/wiki/':
            page = self.wiki2doc.get_wikipage(hyper[2])
            if not page:
                self.errorlog_missing_page(hyper[2], spec_name)
            hyperlink = get_base_url(self.req) +\
                remove_forward_slash(hyper[1]) +\
                hyper[2]
        elif hyper[1] == 'e:/wiki/':
            page = self.wiki2doc.get_wikipage(hyper[2])
            if not page:
                self.errorlog_missing_page(hyper[2], spec_name)
            hyperlink = get_base_url(self.req) +\
            'wiki/' + hyper[2]
        elif hyper[1] == 'wiki:':
            page = \
                self.wiki2doc.get_wikipage(remove_forward_slash(hyper[2]))
            if not page:
                self.errorlog_missing_page(hyper[2], spec_name)
            hyperlink = get_base_url(self.req) +\
                'wiki/' + hyper[2]
        else:
            hyperlink = hyper[1] + hyper[2]
  
        return hyperlink
  
    def get_hyperlink(self, spec_name, regex_id, hyper):
        """ for a given hypermatch this function
            returns the hyperlink and the link name."""
  
        link_name = get_link_name(hyper)
        hyperlink = self.get_link_path(spec_name, regex_id, hyper)
  
        return (hyperlink, link_name)
  
    def filter_hyperlinks(self, args):
        """ for a given paragraph text or a table text,
            this function filters the table text
            and returns the table data
            args = [table,
                    paragraph,
                    text,
                    spec_name]"""
  
        #context = Context.from_request(self.req, 'wiki')
        context = web_context(self.req, 'wiki')
        #print('inside filter_hyperlinks -> everthing is ok so far before -> regex_id, hypermatches = find_hyperlinks(args[2])')
        regex_id, hypermatches = find_hyperlinks(args[2])
        #print('regex_id, hypermatches = find_hyperlinks(args[2])', regex_id, hypermatches)
        hyperlink = ''
        if len(hypermatches) > 0:
                link_name = ''
                rest = hypermatches.pop()
                for hyper in hypermatches:
                    flt_text = ''
                    if self.add_hyper_link == True:
                        #args[2].rows[args[0]].cells[args[1]].\
                        #    paragraphs[0].add_run(hyper[0])
                        wiki = process_blockquote(check_string(hyper[0]))
                        self.parse_html(args, context, wiki)
                        hyperlink, link_name = self.get_hyperlink(args[3],
                                                                  regex_id,
                                                                  hyper)
                        if hyperlink == None:
                            break
                        add_hyperlink(args[1],
                                      hyperlink,
                                      link_name,
                                      '0000FF',
                                      True)
                    elif self.add_hyper_link == False:
                        flt_text = flt_text + hyper[0]
                        #args[2].rows[args[0]].cells[args[1]].\
                        #    paragraphs[0].add_run(flt_text)
                        wiki = process_blockquote(check_string(flt_text))
                        self.parse_html(args, context, wiki)
                #args[2].rows[args[0]].cells[args[1]].paragraphs[0].add_run(rest)
                wiki = process_blockquote(check_string(rest))
                self.parse_html(args, context, wiki)
        else:
            #print('inside filter_hyperlinks(self, args): after else:')
            #print('args[2]', args[2])
            wiki = process_blockquote(check_string(args[2]))
            #print('wiki', wiki)
            #print('args', args)
            self.parse_html(args, context, wiki)
            #print('fine so far after self.parse_html(args, context, wiki) in inside filter_hyperlinks(self, args): after else:')
            #args[2].rows[args[0]].cells[args[1]].text = \
            #        unicode(args[3], "utf-8")
        return (args[0], hypermatches)
  
    def parse_html(self, args, context, wiki):
        """ Parse html string to docx
        args[1] = paragraph,
        context,
        wiki,
        args[3] = spec_name"""
  
        try:
            html_code = HtmlFormatter(self.env,
                                      context,
                                      wiki).generate()
            #print('inside parse_html before DocumentHTMLParser')
            DocumentHTMLParser(self.document, args[1], html_code)
            #print('inside parse_html after DocumentHTMLParser')
            return html_code
        except AttributeError:
            self.wiki2doc.errorlog.append(
                ("HtmlFormatter could not parse" +\
                 " the following wikitext: {}".format(wiki),
                 get_base_url(self.req) + 'wiki/' + args[3]))
  
    def find_sections(self, params):
        """ Given paragraph location and sections data,
            inserts section text, if found images and
            if found tables.
            params = [i,
                      paragraph,
                      sections,
                      text,
                      spec_images]"""
        img_filename = None
        img_path = None
        wiki_filter = \
            [re.compile(r'\s*\[\[Image\((.*(\.jpg|\.png|\.gif))\)\]\]\s*'),
             re.compile(r'\s*\[\[Table\((.*)\.tbl\)\]\]\s*'),
             re.compile(r'\s*(=+)(.*)'),
             re.compile(r'\s*\[\s*=\#Table(\d+)\s*\]\s*'),
             re.compile(r'\s*\[\s*=\#Fig(\d+)\s*\]\s*'),
             re.compile(r'\s*\*\s*(.*)')]
#         image = re.compile(r'\s*\[\[Image\((.*(\.jpg|\.png|\.gif))\)\]\]\s*')
#         anchor = re.compile(r'\s*\[\[Table\((.*)\.tbl\)\]\]\s*')
#         section = re.compile(r'\s*(=+)\s*(\d+\.){1,}\d*(.*)')
#         tbl = re.compile(r'\s*\[\s*=\#Table(\d+)\s*\]\s*')
#         fig = re.compile(r'\s*\[\s*=\#Fig(\d+)\s*\]\s*')
        print('find_sections:')
        print('find_sections -> TEXT:', params[3])
        for line in params[3].splitlines():
            line = to_unicode(line)
#             img_match = wiki_filter[0].match(line)
#             anc_match = wiki_filter[1].match(line)
#             sec_match = wiki_filter[2].match(line)
#             tbl_match = wiki_filter[3].match(line)
#             fig_match = wiki_filter[4].match(line)

            get_header_in_text_line(line)
            
            if wiki_filter[0].match(line):
                #print('1')
                img_filename = to_unicode(wiki_filter[0].match(line).group(1))
                for key, value in params[4].iteritems():
                    if key == img_filename:
                        img_path = value
                        # if you want to include the image name
                        # insert the code below
                        # params[1].insert_paragraph_before(line)
                        insert_image(params[1], img_path)
            elif wiki_filter[1].match(line):
                #print('2')
                self.get_table(params,
                               to_unicode(wiki_filter[1].match(line).group(1)))
            elif wiki_filter[2].match(line):
                print('3')
                
                #line = get_header_in_text_line(line)
                print('wiki_filter[2].match(line).group(1)', wiki_filter[2].match(line).group(1))
                print('wiki_filter[2].match(line).group(2)', wiki_filter[2].match(line).group(2))
                style_key = 'Heading' +\
                            ' ' + \
                            str(len(wiki_filter[2].match(line).group(1)))

                print('WHY!')
                print(dir(params[1]),'\n')
                print('\n')
                print(dir(params[1]._p), '\n')
                print('\n')
                print(dir(params[1]._p.add_p_before), '\n')
                print('\n')
                print(dir(params[1]._insert_paragraph_before), '\n')
                print('\n')
                print(dir(params[1]._p.add_p_before()))
                print('\n')
                
                params[1].insert_paragraph_before(\
                    to_unicode(wiki_filter[2].match(line).group(2).strip()),
                    style=style_key)
                
                #params[1] = self.insert_paragraph_after(params[1], "Paragraph One And A Half.")
#                 if params[1] is not None:
#                     new = params[1].insert_paragraph_before()
#                     new.text = 'test'
#                     p = new._p
#                     p.addnext(new._p)
                    

                if params[1] is not None:
                    new = params[1].insert_paragraph_before(to_unicode(wiki_filter[2].match(line).group(2).strip()), style=style_key)
                    #new.text = 'test'
                    run = new.add_run('test')
                    run.font.subscript = True
                    p = new._p
                    p.addnext(new._p)                    
                
            elif wiki_filter[3].match(line):
                #print('4')
                line = 'Table' + ' ' + str(wiki_filter[3].match(line).group(1))
                line = to_unicode(line)
                params[1].insert_paragraph_before(line, style='Caption')
            elif wiki_filter[4].match(line):
                #print('5')
                line = 'Figure' + ' ' + str(wiki_filter[4].match(line).group(1))
                line = to_unicode(line)
                params[1].insert_paragraph_before(line, style='Caption')
            elif wiki_filter[5].match(line):
                #print('6')
                line = str(wiki_filter[5].match(line).group(1))
                line = to_unicode(line)
                paragraph = create_list(\
                    params[1].insert_paragraph_before(text=' ',
                                                      style='List Bullet'))
                line = filter_wiki_text(line)
                args = [None,
                        paragraph,
                        line,
                        params[2][params[0]][0]]
                # spec_name -> params[2][params[0]][0] = sections[0]
                self.filter_hyperlinks(args)
            else:
                #print('7')
                #print('line:', line)
                #print('params[2][params[0]][0]:', params[2][params[0]][0])
                #print('params[2][params[0]][1]:', params[2][params[0]][1])
                
                line = filter_wiki_text(line)
                args = [None,
                        params[1].insert_paragraph_before(),
                        line,
                        params[2][params[0]][0]]
                # spec_name -> params[2][params[0]][0] = sections[0]
                self.filter_hyperlinks(args)
  
    def get_table(self, params, match_group):
        """ Gets table information from
            list of sections and calls
            insert_table method.
            example_sections = [
            [2,
             'Specname1',
             'consetetur sadipscing elitr, sed diam nonumy eirmod tempor\n' +\
             '[[Image(Image1.jpg)]]\n' +\
             'invidunt ut labore et dolore magna \n' +\
             'aliquyam erat, sed diam \n' +\
             '[[Image(Image2.jpg)]]\nvoluptua.\n',
             {'Image1.jpg': u'/tmp/trac-tempenv-WbQieJ/files/attachments/' +\
              'wiki/bdc/bdc726f49cd502d4306404b090a5ddd13bb7dc0e/98c78c01' +\
              'ccdb21a78fd4f561e980ccd4d3a5a685.jpg',
              'Image2.jpg': u'/tmp/trac-tempenv-WbQieJ/files/attachments/' +\
              'wiki/bdc/bdc726f49cd502d4306404b090a5ddd13bb7dc0e/e8385af6' +\
              'dfec928ba93ae7b6ccdc2c5f2fcb89f8.jpg'},
             {'Table_11': [[[' 1 '], [' 2 '], [' 3 ']],
                           [[''], [' 1-2 '], [' 3 ']],
                           [[' 1 '], [''], [' 2-3 ']],
                           [['', ''], [' 1-2-3 ']]]}],
               ...]]
            params = [i,
                      paragraph,
                      sections,
                      text,
                      spec_images]"""
  
        if params[2][params[0]][3]:
            for value in params[2][params[0]][3]:
                if value == match_group:
                    table_data = params[2][params[0]][3][value]
                    
                    print('params[1]:', params[1])
                    print('table_data:', table_data)
                    print('params[2][params[0]][0]:', params[2][params[0]][0])
                    print('params[2][params[0]][1]:', params[2][params[0]][1])
                       
                    self.insert_table(params[1],
                                      table_data,
                                      params[2][params[0]][0])
  
    def insert_analysed_apos_table(self, paragraph, sections):
        """ Given paragraph location and sections data,
            creates a table that contains analysed
            apo information."""
  
        col_names = ('ITEM', 'ANALYSE APO TASK No', 'APO TASK NAME', 'REMARK')
  
        table = self.document.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER # pylint: disable=no-member
        table.style = 'Table Grid'
  
        for idx, name in enumerate(col_names):
            prg = table.rows[0].cells[idx].paragraphs[0]
            run = prg.add_run(name)
            run.font.name = 'Arial Black'
            run.bold = True
            run.italic = True
  
        for idx in range(len(sections)):
            table.add_row()
            prg = table.rows[idx+1].cells[0].paragraphs[0]
            run = prg.add_run(str(idx+1))
            run.bold = True
            for idy, cell in enumerate(table.rows[idx+1].cells):
                if idy == 1:
                    cell.paragraphs[0].add_run(str(sections[idx][0]))
                elif idy == 2:
                    cell.paragraphs[0].add_run(str(sections[idx][1]))
  
        table = table_font_size(table, 8)
  
        if paragraph is not None:
            new = paragraph.insert_paragraph_before()
            # _p and _tbl are protected and therefore not documented,
            # but otherwise it is impossible to insert a table,
            # and it would only be possible to append it to the
            # end of the document
            new._p.addnext(table._tbl) # pylint: disable=protected-access
  
    def insert_section(self, paragraph, sections, level):
        """ Given paragraph location and sections data,
            inserts section text, if found images and
            if found tables."""

        spec_images = {}
        print('inside_section -> sections', sections)
        for i in range(len(sections)):
            print('inser_section -> TEXT:', i, sections[i][1])
            text = sections[i][1]
            print(text)
            spec_images.update(sections[i][2])
            style_key = 'Heading '+ str(level)
            #paragraph.insert_paragraph_before(apo_spec, style=style_key)
            if text is not None:
                params = [i, paragraph, sections, text, spec_images]
                print('Inside insert_section params=', params)
                self.find_sections(params)

    def add_document(self, sections):
        """ adds intoduction section/s in the spec/s to
            the introduction section in the document."""

        paragraph = self.get_paragraph_after_regex(r"")
        print('inside add_document after paragraph = self.get OK SO FAR!')

        self.insert_section(paragraph, sections, 3)
